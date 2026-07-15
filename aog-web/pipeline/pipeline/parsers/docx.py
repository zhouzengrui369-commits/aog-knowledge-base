"""DOCX 解析: 段落 + 表格 → markdown 文本 + 结构化数据。

AOG 知识库的 .docx 有两类:
- 城市预案: 1 个大表, 分 section (机场信息 / 吉祥执飞 / 航材保障预案 / ...)
- 经验/核心预案: 段落式 (Title/Heading 1/Heading 2/Normal), 无表

本 parser 同时输出两者, 供 extractors 按需取用。
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Union

from docx import Document
from docx.document import Document as _Document

PathLike = Union[str, Path]


# AOG 知识库 city docx 的 section 标签。cell[0] 是 section 标签时, 进入新 section。
# 其它行 (cell[0] 是数据 org/型号/件号 等) 累积到当前 section。
SECTION_LABELS = {
    "机场信息",
    "吉祥执飞",
    "航材保障预案",
    "当地及周边资源",
    "仓储单位",
    "营业部",
    "物流运输",
    "进出口报关",
    "联系人",
    "备件清单",
    "执飞",
    "保障预案",
    "周边资源",
    "仓储",
    "运输",
    "物流",
}


def _is_section_label(label: str) -> bool:
    """判断 cell[0] 是否是 section 标签。支持 prefix / contains 匹配。

    e.g. '吉祥执飞 航班类型' 应该匹配 '吉祥执飞'
         '航材保障预案' 严格匹配
    """
    if not label:
        return False
    label = label.strip()
    for sl in SECTION_LABELS:
        if label == sl:
            return True
        # 双向包含 (cell[0] 含 label, 或 label 含 cell[0]) - 用于 '吉祥执飞 航班类型'
        if sl and (sl in label or label.startswith(sl)):
            return True
    return False


@dataclass
class DocxParagraph:
    """docx 中一个段落, 带 style 名称。"""

    style: str
    text: str


@dataclass
class DocxSection:
    """docx 中一个表格 section: name + 多行 row 数据。"""

    name: str
    rows: list[list[str]] = field(default_factory=list)

    def is_separator(self) -> bool:
        return not self.name.strip()


@dataclass
class DocxTable:
    """docx 顶层解析结果: 段落 + 表格 sections。"""

    paragraphs: list[DocxParagraph] = field(default_factory=list)
    title_rows: list[list[str]] = field(default_factory=list)  # 跨列表标题 (e.g. 机场名)
    sections: list[DocxSection] = field(default_factory=list)  # 表格 sections (城市预案类)


def _normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def parse_docx(path: PathLike) -> DocxTable:
    """读 .docx → DocxTable 结构化对象。

    解析:
    1. 所有段落 (style + text)
    2. 如果有表格, 解析为 sections (按 cell[0] 分组)
    3. 跨列同名行 → title_rows (e.g. 城市名)
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"docx 文件不存在: {p}")
    if p.suffix.lower() != ".docx":
        raise ValueError(f"不是 .docx 文件: {p}")

    doc: _Document = Document(str(p))
    result = DocxTable()

    # 段落
    for para in doc.paragraphs:
        text = _normalize(para.text)
        if not text and (para.style is None or "Title" not in (para.style.name or "")):
            continue
        style_name = para.style.name if para.style else "Normal"
        result.paragraphs.append(DocxParagraph(style=style_name, text=text))

    # 表格 (取第一个; AOG 知识库每个 docx 只有一个)
    if not doc.tables:
        return result

    table = doc.tables[0]
    rows = list(table.rows)
    current_section: DocxSection | None = None
    seen_labels: set[str] = set()  # 已经开过的 section label

    for row in rows:
        cells = [_normalize(c.text) for c in row.cells]
        # 标题行 (跨列同名) - 如果整行所有 cell 文本一致且非空, 是标题
        if cells and all(c == cells[0] for c in cells) and cells[0]:
            result.title_rows.append(cells)
            continue
        label = cells[0] if cells else ""
        # 决定新 section:
        # 1) cell[0] 是 SECTION_LABELS 之一 (含 prefix 匹配) → 新 section
        # 2) cell[0] 是空 → 跳过分隔行
        # 3) 其它 (e.g. "东航", "A320", "前轮") → 数据行, 累积到当前 section
        is_section_label = _is_section_label(label)
        if is_section_label and label not in seen_labels:
            # 新 section
            current_section = DocxSection(name=label)
            result.sections.append(current_section)
            seen_labels.add(label)
            current_section.rows.append(cells)
        elif is_section_label:
            # 重复的 section label (cell[0] 还是这个 section 标签) → 累积
            for sec in reversed(result.sections):
                if sec.name == label:
                    sec.rows.append(cells)
                    current_section = sec
                    break
        elif not label:
            continue
        else:
            if current_section is None:
                current_section = DocxSection(name="")
                result.sections.append(current_section)
            current_section.rows.append(cells)

    return result


def _para_to_markdown(p: DocxParagraph) -> str:
    """段落 → markdown 单行。"""
    s = (p.style or "").lower()
    t = p.text
    if "title" in s:
        return f"# {t}"
    if "heading 1" in s or s == "heading 1":
        return f"## {t}"
    if "heading 2" in s or s == "heading 2":
        return f"### {t}"
    if "heading 3" in s or s == "heading 3":
        return f"#### {t}"
    return t


def docx_table_to_markdown(dt: DocxTable) -> str:
    """把 DocxTable 序列化为 markdown, 喂给 chroma 索引用。

    城市预案类 (有 sections):
        # 标题 (R0)
        ## section_name
        | col1 | col2 | ... |
        | --- | --- | --- |
        | val1 | val2 | ... |

    段落式 (B787 风挡 AOG 经验):
        # Title
        ## Heading 1
        content
    """
    lines: list[str] = []

    # 1. 段落 (有 style 的)
    for p in dt.paragraphs:
        md = _para_to_markdown(p)
        if md:
            lines.append(md)
            lines.append("")

    # 2. title rows (跨列表标题, e.g. 机场名)
    for title_row in dt.title_rows:
        if title_row and title_row[0] and not lines:
            # 避免与段落 # Title 重复
            lines.append(f"# {title_row[0]}")
            lines.append("")

    # 3. 表格 sections
    for sec in dt.sections:
        if not sec.name:
            continue
        # 如果 lines 已经有 section 标题, 跳过
        if not any(l.strip() == f"## {sec.name}" for l in lines):
            lines.append(f"## {sec.name}")
            lines.append("")

        if not sec.rows:
            continue
        header = sec.rows[0]
        seen: list[str] = []
        for h in header:
            if h and (not seen or seen[-1] != h):
                seen.append(h)
        if not seen:
            continue

        # 表头
        lines.append("| " + " | ".join(seen) + " |")
        lines.append("| " + " | ".join(["---"] * len(seen)) + " |")

        for row in sec.rows[1:]:
            if all(not c for c in row):
                continue
            values = row[: len(seen)]
            while len(values) < len(seen):
                values.append("")
            lines.append("| " + " | ".join(v.replace("|", "\\|") for v in values) + " |")

        lines.append("")

    return "\n".join(lines).strip()


def parse_docx_to_markdown(path: PathLike) -> str:
    """便捷: 读 .docx 直接返回 markdown 字符串。"""
    dt = parse_docx(path)
    return docx_table_to_markdown(dt)
