#!/usr/bin/env python3
"""wiki_curator.py — LLM 周期整理 AOG 知识库为结构化 wiki 页面 (NJX 14:43 拍 🅰️ 双轨方案)

🅰️ 双轨架构:
  - 后台 (本脚本): 扫 docx → 调 MiniMax M3 → 生成 MOC 风格 wiki 页面
  - 前台: 用户浏览 wiki + AI chat RAG (wiki + docx 双重召回)

输入:
  - /Users/njx/Project/AOG知识库/AOG知识库/02_外战预案/*.docx (城市预案, 主流)
  - /Users/njx/Project/AOG知识库/AOG知识库/01_AOG预案/*.{md,xlsx} (核心预案)
  - /Users/njx/Project/AOG知识库/AOG知识库/03_保障经验/*.{md,docx,xlsx} (经验)

输出 (staging, 不污染 read-only 源):
  - pipeline/data/wiki/MOC-{city_code}-{topic}.md (如 MOC-X-西安-故障树.md)
  - pipeline/data/wiki/.meta.json (chunk_id 映射 + docx 引用溯源)
  - pipeline/data/wiki/INDEX.md (wiki 总目录)

NSM-2 红线:
  - 每 wiki 页面 强制末尾列出 "## 引用" 段, 指向源 docx 路径
  - LLM 输出含未确认信息时必须打 "⚠️ 需 NJX 核实" 标记
  - 不删源 docx, 只生成新 wiki 页面 (用户可 cp 到 AOG知识库/00_MOC/ 评审)

LLM:
  - 复用 backend minimax (从 backend/.env 读 MINIMAX_API_KEY/MODEL/BASE_URL)
  - 单 docx 一调, 温度 0.3 (结构化 + 不幻觉)
  - max_tokens 4000 (一个 MOC 页通常 1500-2500 token)

使用:
  # 跑所有 225 城市 (full)
  uv run python -m scripts.wiki_curator

  # MVP: 只跑 3 个示范城市 (NJX 评审用)
  uv run python -m scripts.wiki_curator --codes X-西安 B-北京大兴 S-三亚

  # 跑特定 topic (如只生成故障树)
  uv run python -m scripts.wiki_curator --topic 故障树

  # dry-run (不调 LLM, 只列计划)
  uv run python -m scripts.wiki_curator --dry-run
"""
from __future__ import annotations

import argparse
import json
import logging
import os
import re
import sys
import time
from dataclasses import asdict, dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

# 路径: 复用 backend config (从 backend/.env 读 minimax key)
PIPELINE_DIR = Path(__file__).resolve().parent.parent
WORKSPACE = PIPELINE_DIR.parent
sys.path.insert(0, str(WORKSPACE / "backend"))

# ⚠️ read-only 源目录 (NEVER modify)
KB_ROOT = Path("/Users/njx/Project/AOG知识库/AOG知识库")
CITIES_DIR = KB_ROOT / "02_外战预案"
CORE_PLANS_DIR = KB_ROOT / "01_AOG预案"
EXPERIENCES_DIR = KB_ROOT / "03_保障经验"

# 输出 (staging, 不污染源)
WIKI_OUT_DIR = PIPELINE_DIR / "data" / "wiki"
WIKI_META_PATH = WIKI_OUT_DIR / ".meta.json"
WIKI_INDEX_PATH = WIKI_OUT_DIR / "INDEX.md"

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger("wiki_curator")


# ====== docx 抽取 ======

def extract_docx(path: Path) -> str:
    """从 docx 抽 text + table 内容"""
    from docx import Document
    d = Document(str(path))
    parts: List[str] = []
    # paragraph
    for p in d.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    # table (按行)
    for tbl in d.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_md(path: Path) -> str:
    """md 全文"""
    return path.read_text(encoding="utf-8", errors="ignore")


def extract_excel(path: Path) -> str:
    """xlsx 抽 (取所有 sheet 前 50 行)"""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(path, data_only=True, read_only=True)
        parts: List[str] = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"# Sheet: {sheet_name}")
            for row in ws.iter_rows(max_row=50, values_only=True):
                if any(c is not None for c in row):
                    parts.append(" | ".join(str(c) if c is not None else "" for c in row))
        return "\n".join(parts)
    except ImportError:
        return ""


def extract_any(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx(path)
    if suffix in {".md", ".markdown"}:
        return extract_md(path)
    if suffix in {".xlsx", ".xls"}:
        return extract_excel(path)
    return ""


# ====== LLM 调用 ======

def call_llm(prompt: str, system: str = "", max_tokens: int = 12000, temperature: float = 0.3) -> str:
    """调 minimax M3, 复用 backend llm 客户端"""
    try:
        # 复用 backend llm 服务 (从 backend/.env 读 env)
        from aog_web.services.llm import get_llm
        from aog_web.config import get_settings
        s = get_settings()
        llm = get_llm(settings=s)
        messages = []
        if system:
            messages.append({"role": "system", "content": system})
        messages.append({"role": "user", "content": prompt})
        # 同步调用 (不用 async, curator 是后台任务)
        # ⚠️ 关键: 必须传 max_tokens + temperature 给 llm.chat()
        #   backend minimax.py default max_tokens=1024 (只够 think 段, wiki 截断)
        import asyncio
        return asyncio.run(llm.chat(messages, max_tokens=max_tokens, temperature=temperature))
    except Exception as e:
        logger.error("LLM call failed: %s", e)
        return f"⚠️ LLM 调用失败: {e}"


# ====== prompt 模板 ======

SYSTEM_PROMPT = """你是 AOG（飞机停场维修）知识库的 wiki 整理员。你的任务是把原始 docx 整理成结构化 wiki 页面。

⚠️ 重要: 你的输出必须**只包含** sentinel 之间的内容, 不要输出任何思考过程 / 解释 / 元评论。

要求:
1. **NSM-2 红线**: 每页必须末尾有 "## 引用" 段, 列出源 docx 路径, 任何不确定信息打 "⚠️ 需 NJX 核实"
2. **MOC 风格**: 用 markdown, 标题 ## ##, 故障树用 - 列表缩进
3. **保留原文术语**: 件号 (3-1531-3), IATA (PVG), 航司名 (东航/上航), 联系人姓名电话
4. **不要编造**: 原文没说的人/电话/件号不要补, 没数据写 "无原文"
5. **交叉链接**: 提到其他城市时用 [[S-上海浦东]] wiki link 格式
"""


def build_prompt(city_code: str, city_name: str, docx_text: str, topic: str) -> str:
    """构造整理 prompt (用 sentinel 让 LLM 严格输出 wiki)"""
    return f"""# 任务
把下面的 AOG 城市预案 docx 整理成 wiki 页面 (主题: {topic})。

# 城市
- 代码: {city_code}
- 名称: {city_name}

# 源 docx 原文
```
{docx_text[:6000]}  # 截断 6000 字符 (避免超 token)
```

# 输出格式 (严格遵守, 只输出 sentinel 之间的内容)
===WIKI_START===
# {city_name} ({city_code}) — {topic}

## 基本信息
- IATA: (从原文)
- 省份/地区: (从原文)
- 状态: (现行/暂停/已废)

## 故障树 / 决策表 / 备件清单 / 联系方式
(按主题分类, 原文表格/列表转 markdown)

## 互援关系
(从原文"当地及周边资源"提取, 用 [[code]] 标记其他城市)
- 互援: [[X-城市]] — 原因
- 中介: [[X-城市]] — 原因

## 风险与备注
(从原文提取, 不确定信息打 ⚠️)

## 引用
- 源: 02_外战预案/{city_code}.docx
- 整理时间: (LLM 生成时间)
- 整理者: MiniMax M3 (minimax-m3)
===WIKI_END===

⚠️ 任何思考/解释/元评论请放在 ===WIKI_END=== 之后 (或省略), 严禁放在 ===WIKI_START=== 之前"""


# Wiki 内容 sentinel 标记 (让 LLM 严格输出 wiki, post-process 切出来)
WIKI_SENTINEL_START = "===WIKI_START==="
WIKI_SENTINEL_END = "===WIKI_END==="


def extract_wiki_from_response(response: str) -> str:
    """从 LLM 响应切出 wiki 内容 (剥掉 think 段 + 解释)"""
    if WIKI_SENTINEL_START in response and WIKI_SENTINEL_END in response:
        start = response.index(WIKI_SENTINEL_START) + len(WIKI_SENTINEL_START)
        end = response.index(WIKI_SENTINEL_END)
        return response[start:end].strip()
    # fallback: 剥掉 <think>...</think> 段
    if "<think>" in response and "</think>" in response:
        end = response.index("</think>") + len("</think>")
        return response[end:].strip()
    return response.strip()


# ====== 整理主流程 ======

@dataclass
class WikiPage:
    code: str
    name: str
    topic: str
    title: str
    content: str
    source_path: str
    generated_at: str
    llm_model: str
    char_count: int = 0


def curate_one(code: str, name: str, docx_path: Path, topic: str = "故障树") -> Optional[WikiPage]:
    """整理一个 docx → wiki page"""
    logger.info("curate: %s (%s), topic=%s", code, name, topic)
    text = extract_any(docx_path)
    if not text:
        logger.warning("extract failed: %s", docx_path)
        return None

    prompt = build_prompt(code, name, text, topic)
    # P0 治本 (NJX 7/27 wiki_curator X-西安): max_tokens 8000 截断 45 chars
    # X-西安 docx 3047 chars + think 段 = 8000 token 不够, 改 12000
    raw = call_llm(prompt, system=SYSTEM_PROMPT, max_tokens=12000, temperature=0.3)
    if raw.startswith("⚠️"):
        logger.error("LLM fail for %s: %s", code, raw[:100])
        return None
    # 切出 wiki 内容 (剥掉 think 段 + 解释)
    content = extract_wiki_from_response(raw)
    if not content or len(content) < 100:
        logger.warning("wiki content too short for %s: %d chars", code, len(content))
        return None

    now = datetime.now(timezone.utc).isoformat()
    title = f"{name} ({code}) — {topic}"
    return WikiPage(
        code=code,
        name=name,
        topic=topic,
        title=title,
        content=content,
        source_path=str(docx_path.relative_to(KB_ROOT.parent)),
        generated_at=now,
        llm_model="minimax-m3",
        char_count=len(content),
    )


def write_wiki_page(page: WikiPage, out_dir: Path = WIKI_OUT_DIR) -> Path:
    """写 wiki 页面到 staging 目录"""
    out_dir.mkdir(parents=True, exist_ok=True)
    fname = f"MOC-{page.code}-{page.topic}.md"
    out_path = out_dir / fname
    frontmatter = f"""---
code: {page.code}
name: {page.name}
topic: {page.topic}
source: {page.source_path}
generated_at: {page.generated_at}
llm: {page.llm_model}
chars: {page.char_count}
---

"""
    out_path.write_text(frontmatter + page.content, encoding="utf-8")
    logger.info("wrote %s (%d chars)", out_path, page.char_count)
    return out_path


def update_meta(pages: List[WikiPage]) -> None:
    """更新 .meta.json"""
    WIKI_OUT_DIR.mkdir(parents=True, exist_ok=True)
    meta = {
        "updated_at": datetime.now(timezone.utc).isoformat(),
        "total_pages": len(pages),
        "pages": [asdict(p) for p in pages],
    }
    WIKI_META_PATH.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")


def write_index(pages: List[WikiPage]) -> None:
    """写 wiki 总目录 INDEX.md"""
    WIKI_OUT_DIR.mkdir(parents=True, exist_ok=True)
    by_topic: Dict[str, List[WikiPage]] = {}
    for p in pages:
        by_topic.setdefault(p.topic, []).append(p)
    lines: List[str] = ["# AOG Wiki 总目录\n", f"_Generated at {datetime.now(timezone.utc).isoformat()}_\n", ""]
    for topic, ps in sorted(by_topic.items()):
        lines.append(f"## {topic} ({len(ps)} 篇)\n")
        for p in sorted(ps, key=lambda x: x.code):
            rel = f"MOC-{p.code}-{p.topic}.md"
            lines.append(f"- [{p.title}](./{rel}) — 引用: `{p.source_path}`")
        lines.append("")
    WIKI_INDEX_PATH.write_text("\n".join(lines), encoding="utf-8")
    logger.info("wrote INDEX.md (%d topics, %d pages)", len(by_topic), len(pages))


def main():
    parser = argparse.ArgumentParser(description="wiki_curator: LLM 整理 AOG 知识库")
    parser.add_argument("--codes", nargs="*", help="只跑这些 city code (e.g. X-西安 B-北京大兴)")
    parser.add_argument("--topic", default="故障树", help="wiki 主题 (故障树/决策表/备件清单...)")
    parser.add_argument("--dry-run", action="store_true", help="只列计划不调 LLM")
    args = parser.parse_args()

    # 扫 city docx
    city_files = sorted(CITIES_DIR.glob("*.docx"))
    if args.codes:
        wanted = set(args.codes)
        city_files = [f for f in city_files if f.stem in wanted]
    logger.info("扫描 %d 个城市 docx (codes=%s)", len(city_files), args.codes or "ALL")

    pages: List[WikiPage] = []
    for docx in city_files:
        code = docx.stem
        name = code.split("-", 1)[1] if "-" in code else code
        if args.dry_run:
            print(f"[DRY] would curate: {code} ({name})")
            continue
        t0 = time.time()
        page = curate_one(code, name, docx, topic=args.topic)
        if page:
            write_wiki_page(page)
            pages.append(page)
        logger.info("done %s in %.1fs", code, time.time() - t0)

    if pages:
        update_meta(pages)
        write_index(pages)
        print(f"\n✓ Generated {len(pages)} wiki pages at {WIKI_OUT_DIR}")
        print(f"  Index: {WIKI_INDEX_PATH}")
        print(f"  Meta:  {WIKI_META_PATH}")
    elif args.dry_run:
        print(f"\n[DRY] would process {len(city_files)} cities")
    else:
        print("\n✗ No pages generated")


if __name__ == "__main__":
    main()
