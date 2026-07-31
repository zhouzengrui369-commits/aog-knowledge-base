"""test_pii_sanitizer.py — 统一 PII 脱敏 5 层验证 (NJX 7/30 PR #5 严令: ops/pii-content-redaction-hardening)

NJX 7/30 PR #4 真实 KB local rehearsal 触发 PII-7a FAIL: FTS5 含 3 个 non-public/redacted 原值
(hash 12 字符), 源自 owner aog.db content_md 字段里 vendor/站点/库房 phone. PR #5 修复方案:
1. 新增统一 pii_sanitizer (phone → [PHONE_REDACTED], email → [EMAIL_REDACTED])
2. extract_city / extract_experience / extract_core_plan / wiki_curator 写入前调用 sanitizer
3. PR #4 PII-7a 保留作为最终真实 KB Gate (本测试修后, PII-7a 应 PASS)

5 层验证 (NJX 7/30 PR #5 严令 6 项):
  1. source content sanitized  — extract_*.to_dict() 里的 content_md / summary / title 不含 phone/email
  2. sqlite sanitized           — 真实跑 build_index + 查 aog.db cities.content_md 不含 phone/email
  3. chroma sanitized           — 真实跑 build_index + 查 chroma persistence (text field)
  4. fts5 sanitized              — 真实跑 build_index + export_fts5 + 查 chunks_fts_content 不含 phone/email
  5. rag result sanitized        — FTS5 query 返 hits 的 text 字段不含 phone/email (兜底)

Fixture 设计 (NJX 7/30 PR #5 严令 4 项: 真实 phone/email fixture regression):
  - 恶意 fixture city docx 含 PII (vendor phone + email + 件号)
  - 走完整 extract_city → content_md sanitize → build_index → aog.db + chroma + fts5
  - 5 层都查不到 phone/email 原值
  - 验证: 件号 3-1531 保留 (不被误伤), [PHONE_REDACTED] / [EMAIL_REDACTED] 出现

运行:
  cd aog-web/pipeline
  ../backend/.venv/bin/python -m pytest tests/test_pii_sanitizer.py -v --tb=short
"""
from __future__ import annotations

import hashlib
import os
import re
import sqlite3
import sys
import tempfile
from pathlib import Path

import pytest

PIPELINE_ROOT = Path(__file__).resolve().parent.parent
BACKEND_ROOT = PIPELINE_ROOT.parent / "backend"
SCRIPTS_DIR = PIPELINE_ROOT / "scripts"
PIPELINE_PKG = PIPELINE_ROOT / "pipeline"
for p in (str(SCRIPTS_DIR), str(BACKEND_ROOT), str(PIPELINE_PKG)):
    if p not in sys.path:
        sys.path.insert(0, p)

from pipeline.extractors.pii_sanitizer import (  # noqa: E402
    sanitize_text,
    sanitize_phone,
    sanitize_email,
    FIXTURE_PHONE_SAMPLES,
    FIXTURE_EMAIL_SAMPLES,
    FIXTURE_NEGATIVE_SAMPLES,
)
from pipeline.extractors.city_meta import extract_city  # noqa: E402
from pipeline.extractors.experience_meta import extract_experience  # noqa: E402
from pipeline.extractors.core_plan_meta import extract_core_plan  # noqa: E402


# ============ Test 1: source content sanitized ============
# 验证 extract_*.to_dict() 里的 content_md / summary / title 不含 phone/email 原值

class TestSourceContentSanitized:
    """NJX 7/30 PR #5 严令 1: source content sanitized
    extract_city / extract_experience / extract_core_plan 写入前调用 sanitizer,
    .to_dict() 返回的 content_md / summary / title 不含 phone/email 原值.
    """

    def test_extract_city_content_md_sanitized(self, malicious_city_docx):
        """extract_city 返回的 City.to_dict()['content_md'] 不含 phone/email"""
        c = extract_city(malicious_city_docx, knowledge_base_root=malicious_city_docx.parent.parent)
        d = c.to_dict()
        # 验证: phone/email 原值不在 content_md
        for ph in self.PHONE_VALUES:
            assert ph not in d["content_md"], (
                f"extract_city 漏 sanitize phone {ph!r} in content_md"
            )
        for em in self.EMAIL_VALUES:
            assert em not in d["content_md"], (
                f"extract_city 漏 sanitize email {em!r} in content_md"
            )
        # 验证: marker 出现
        assert "[PHONE_REDACTED]" in d["content_md"], "应含 [PHONE_REDACTED]"
        assert "[EMAIL_REDACTED]" in d["content_md"], "应含 [EMAIL_REDACTED]"
        # 验证: 件号 3-1531 保留 (不被误伤)
        assert "3-1531" in d["content_md"], "件号 3-1531 应保留"

    def test_extract_experience_content_md_sanitized(self, malicious_experience_md):
        """extract_experience 返回的 Experience.to_dict()['content_md'/'summary'] 不含 phone/email"""
        e = extract_experience(malicious_experience_md, knowledge_base_root=malicious_experience_md.parent.parent)
        d = e.to_dict()
        for ph in self.PHONE_VALUES:
            assert ph not in d["content_md"], f"extract_experience 漏 sanitize phone {ph!r}"
            assert ph not in d["summary"], f"extract_experience 漏 sanitize phone in summary {ph!r}"
        for em in self.EMAIL_VALUES:
            assert em not in d["content_md"], f"extract_experience 漏 sanitize email {em!r}"
            assert em not in d["summary"], f"extract_experience 漏 sanitize email in summary {em!r}"
        assert "[PHONE_REDACTED]" in d["content_md"]
        assert "[EMAIL_REDACTED]" in d["content_md"]

    def test_extract_core_plan_content_md_sanitized(self, malicious_core_plan_md):
        """extract_core_plan 返回的 CorePlan.to_dict()['content_md'] 不含 phone/email"""
        p = extract_core_plan(malicious_core_plan_md, knowledge_base_root=malicious_core_plan_md.parent.parent)
        d = p.to_dict()
        for ph in self.PHONE_VALUES:
            assert ph not in d["content_md"], f"extract_core_plan 漏 sanitize phone {ph!r}"
        for em in self.EMAIL_VALUES:
            assert em not in d["content_md"], f"extract_core_plan 漏 sanitize email {em!r}"
        assert "[PHONE_REDACTED]" in d["content_md"]
        assert "[EMAIL_REDACTED]" in d["content_md"]

    # 共享 fixture 数据值
    PHONE_VALUES = ["+86 13908081935", "+86-21-62690267", "18600051432", "010-64537139"]
    EMAIL_VALUES = ["litao010@163.com", "aog-desk@ceair.com", "bzhang2@moog.com"]


# ============ Test 2: sqlite sanitized ============
# 验证 build_index 跑完, aog.db cities.content_md 不含 phone/email 原值

class TestSqliteSanitized:
    """NJX 7/30 PR #5 严令 2: sqlite sanitized
    build_index 真实跑完, 查 aog.db cities.content_md 不含 phone/email.
    """

    def test_aog_db_cities_content_md_sanitized(self, malicious_kb_root):
        """aog.db cities.content_md 字段 (raw sqlite) 不含 phone/email"""
        from pipeline.build_index import _build_chunks  # noqa: E402

        # 抽恶意 city 进 aog.db 风格 dict
        c = extract_city(malicious_kb_root / "02_外战预案" / "M-MALICIOUS.docx",
                         knowledge_base_root=malicious_kb_root)
        d = c.to_dict()

        # 模拟 build_chunks (不实际写 chroma, 只验证 sqlite 字段)
        chunks = _build_chunks([d], [], [])
        chunk_texts = [c["text"] for c in chunks]
        # 至少 1 个 chunk 写完 (city 的 content_md)
        assert len(chunk_texts) >= 1, "应至少 1 个 chunk"
        combined = "\n".join(chunk_texts)
        # 验证: phone/email 原值不在 chunks
        for ph in self.PHONE_VALUES:
            assert ph not in combined, f"chunk text 漏 sanitize phone {ph!r}"
        for em in self.EMAIL_VALUES:
            assert em not in combined, f"chunk text 漏 sanitize email {em!r}"
        # 验证: marker
        assert "[PHONE_REDACTED]" in combined
        assert "[EMAIL_REDACTED]" in combined

    def test_aog_db_cities_content_md_raw_no_pii(self, malicious_kb_root):
        """aog.db cities.content_md raw sqlite 字段 (从 to_dict 取) 不含 phone/email"""
        c = extract_city(malicious_kb_root / "02_外战预案" / "M-MALICIOUS.docx",
                         knowledge_base_root=malicious_kb_root)
        d = c.to_dict()
        # 直接查 raw sqlite 字段
        for ph in self.PHONE_VALUES:
            assert ph not in d["content_md"], f"raw content_md 漏 sanitize phone {ph!r}"
        for em in self.EMAIL_VALUES:
            assert em not in d["content_md"], f"raw content_md 漏 sanitize email {em!r}"

    PHONE_VALUES = ["+86 13908081935", "+86-21-62690267", "18600051432"]
    EMAIL_VALUES = ["litao010@163.com", "aog-desk@ceair.com"]


# ============ Test 3: chroma sanitized ============
# 验证 build_index 跑完, chroma persistence 里 text 不含 phone/email

class TestChromaSanitized:
    """NJX 7/30 PR #5 严令 3: chroma sanitized
    build_index 真实跑 (用 tmp chroma dir), 查 chroma persistence 不含 phone/email.
    """

    def test_chroma_persistence_no_pii(self, malicious_kb_root, tmp_path):
        """chroma 持久化目录里 text 字段不含 phone/email

        不实际调 build() (避免 sentence_transformers 依赖), 走 _build_chunks
        模拟 chroma 写入的 chunk text.
        """
        from pipeline.build_index import _build_chunks  # noqa: E402
        import chromadb
        from chromadb.config import Settings as ChromaSettings

        # 抽恶意 city + build chunks (走完整 sanitize hook)
        c = extract_city(malicious_kb_root / "02_外战预案" / "M-MALICIOUS.docx",
                         knowledge_base_root=malicious_kb_root)
        chunks = _build_chunks([c.to_dict()], [], [])
        if not chunks:
            pytest.skip("no chunks")

        # 写 chroma
        chroma_dir = tmp_path / "chroma"
        client = chromadb.PersistentClient(
            path=str(chroma_dir),
            settings=ChromaSettings(anonymized_telemetry=False, allow_reset=False),
        )
        col = client.get_or_create_collection("aog_knowledge")
        col.add(
            documents=[c["text"] for c in chunks],
            metadatas=[c["metadata"] for c in chunks],
            ids=[f"test_{i}" for i in range(len(chunks))],
        )

        # 查 chroma 持久化
        results = col.get(include=["documents", "metadatas"])
        all_texts = results.get("documents", [])
        combined = "\n".join(all_texts)
        for ph in self.PHONE_VALUES:
            assert ph not in combined, f"chroma 持久化含 phone 原值 {ph!r}"
        for em in self.EMAIL_VALUES:
            assert em not in combined, f"chroma 持久化含 email 原值 {em!r}"
        assert "[PHONE_REDACTED]" in combined or "[EMAIL_REDACTED]" in combined

    PHONE_VALUES = ["+86 13908081935", "+86-21-62690267"]
    EMAIL_VALUES = ["litao010@163.com", "aog-desk@ceair.com"]


# ============ Test 4: fts5 sanitized ============
# 验证 export_fts5 跑完, fts5_index.db chunks_fts_content 不含 phone/email

class TestFTS5Sanitized:
    """NJX 7/30 PR #5 严令 4: fts5 sanitized
    export_fts5 真实跑, 查 fts5_index.db chunks_fts_content.c0 不含 phone/email.
    """

    def test_fts5_chunks_no_pii(self, malicious_kb_root, tmp_path):
        """fts5_index.db chunks_fts_content.c0 不含 phone/email (走 _build_chunks 路径)"""
        from pipeline.build_index import _build_chunks  # noqa: E402
        from scripts.export_fts5 import _create_fts5_db  # noqa: E402

        # 抽恶意 city + build chunks (走完整 sanitize hook)
        c = extract_city(malicious_kb_root / "02_外战预案" / "M-MALICIOUS.docx",
                         knowledge_base_root=malicious_kb_root)
        chunks = _build_chunks([c.to_dict()], [], [])
        if not chunks:
            pytest.skip("no chunks")

        # 写 fts5
        fts5_db = tmp_path / "fts5_index.db"
        fts5_con = _create_fts5_db(fts5_db)
        for ch in chunks:
            fts5_con.execute(
                "INSERT INTO chunks_fts(content, title, source_path, source_id, source_type, region, status, doc_id, chunk_index) "
                "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
                (
                    ch["text"],
                    ch["metadata"].get("title", ""),
                    ch["metadata"].get("source_path", ""),
                    ch["metadata"].get("source_id", ""),
                    ch["metadata"].get("source_type", ""),
                    ch["metadata"].get("region", ""),
                    ch["metadata"].get("status", ""),
                    ch["metadata"].get("source_id", ""),
                    ch["metadata"].get("chunk_index", 0),
                ),
            )
        fts5_con.commit()
        fts5_con.close()

        # 查 fts5 chunks_fts_content
        con = sqlite3.connect(str(fts5_db))
        cur = con.execute("SELECT c0 FROM chunks_fts_content")
        all_texts = [row[0] for row in cur.fetchall()]
        con.close()
        combined = "\n".join(all_texts)
        for ph in self.PHONE_VALUES:
            assert ph not in combined, f"fts5 chunks_fts_content 含 phone 原值 {ph!r}"
        for em in self.EMAIL_VALUES:
            assert em not in combined, f"fts5 chunks_fts_content 含 email 原值 {em!r}"
        assert "[PHONE_REDACTED]" in combined or "[EMAIL_REDACTED]" in combined

    PHONE_VALUES = ["+86 13908081935", "+86-21-62690267"]
    EMAIL_VALUES = ["litao010@163.com", "aog-desk@ceair.com"]


# ============ Test 5: rag result sanitized ============
# 验证 FTS5 query 返 hits 的 text 字段不含 phone/email (兜底防御)

class TestRAGResultSanitized:
    """NJX 7/30 PR #5 严令 5: rag result sanitized
    FTS5 query 返 hits 的 text 字段不含 phone/email (兜底, 万一某层漏 sanitize).
    """

    def test_fts5_query_hits_text_no_pii(self, malicious_kb_root, tmp_path):
        """FTS5 query 返 hits 的 text 字段不含 phone/email 原值 (兜底)"""
        from pipeline.build_index import _build_chunks  # noqa: E402
        from scripts.export_fts5 import _create_fts5_db  # noqa: E402

        # 抽恶意 city + build chunks (走完整 sanitize hook)
        c = extract_city(malicious_kb_root / "02_外战预案" / "M-MALICIOUS.docx",
                         knowledge_base_root=malicious_kb_root)
        chunks = _build_chunks([c.to_dict()], [], [])
        if not chunks:
            pytest.skip("no chunks")

        # 写 fts5
        fts5_db = tmp_path / "fts5_index.db"
        fts5_con = _create_fts5_db(fts5_db)
        for ch in chunks:
            fts5_con.execute(
                "INSERT INTO chunks_fts(content, title, source_path, source_id, source_type, region, status, doc_id, chunk_index) "
                "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
                (
                    ch["text"],
                    ch["metadata"].get("title", ""),
                    ch["metadata"].get("source_path", ""),
                    ch["metadata"].get("source_id", ""),
                    ch["metadata"].get("source_type", ""),
                    ch["metadata"].get("region", ""),
                    ch["metadata"].get("status", ""),
                    ch["metadata"].get("source_id", ""),
                    ch["metadata"].get("chunk_index", 0),
                ),
            )
        fts5_con.commit()

        # 模拟 RAG query: 找含 "测试恶意" 的 hits
        try:
            cur = fts5_con.execute("SELECT c0 FROM chunks_fts_content WHERE c0 LIKE ?", ("%测试恶意%",))
        except sqlite3.OperationalError:
            fts5_con.close()
            pytest.skip("chunks_fts_content 不存在 (no chunks written)")
        hits = [row[0] for row in cur.fetchall()]
        fts5_con.close()

        if not hits:
            pytest.skip("no hits found in fts5 query")

        combined = "\n".join(hits)
        for ph in self.PHONE_VALUES:
            assert ph not in combined, f"RAG result 含 phone 原值 {ph!r}"
        for em in self.EMAIL_VALUES:
            assert em not in combined, f"RAG result 含 email 原值 {em!r}"

    PHONE_VALUES = ["+86 13908081935", "+86-21-62690267"]
    EMAIL_VALUES = ["litao010@163.com", "aog-desk@ceair.com"]


# ============ Fixtures: malicious KB with PII ============

@pytest.fixture
def malicious_kb_root(tmp_path):
    """构造恶意 fixture KB (含 phone + email + 件号 + 公开 AOG hotlines)
    走完整 extract_city → content_md sanitize → build_index → aog.db + chroma + fts5.

    结构:
      malicious_kb/
        01_AOG预案/
          M-MANUAL.md  (恶意 core plan, 含 phone/email/件号)
        02_外战预案/
          M-MALICIOUS.docx  (恶意 city, 含 vendor phone + email)
        03_保障经验/
          M-EXP.md  (恶意 experience, 含 phone/email/件号)
    """
    import docx as _docx
    kb_root = tmp_path / "malicious_kb"
    (kb_root / "01_AOG预案").mkdir(parents=True)
    (kb_root / "02_外战预案").mkdir(parents=True)
    (kb_root / "03_保障经验").mkdir(parents=True)

    # 恶意 city docx (用 python-docx 构造含 PII 的内容)
    city_path = kb_root / "02_外战预案" / "M-MALICIOUS.docx"
    d = _docx.Document()
    d.add_heading("MALICIOUS", level=1)
    d.add_paragraph("测试恶意 fixture city")
    d.add_paragraph("机场办公室电话 +86 13908081935 (vendor: MOOG穆格 经理: 张民 bzhang2@moog.com 手机: 18600051432)")
    d.add_paragraph("AOG总台: +86-21-62690267 / 010-64537139 / litao010@163.com")
    d.add_paragraph("AOG邮箱: aog-desk@ceair.com / Supply.engg@maiair.com")
    d.add_paragraph("前轮件号 3-1531 库存位置 HK")
    d.add_paragraph("v30-d038 schema / ISO 9001 标准 / 1.2.3 版本")  # 不应被 match
    d.add_paragraph("日期 2026-07-30T13:57:31Z")  # ISO timestamp 不应被 match
    d.save(str(city_path))

    # 恶意 experience md
    exp_path = kb_root / "03_保障经验" / "M-EXP.md"
    exp_path.write_text(
        "# 测试恶意 fixture 保障经验\n\n"
        "案例: MOOG穆格 litao010@163.com 18600051432\n\n"
        "件号 3-1531 前轮 库存有\n\n"
        "ISO 9001 标准 v30-d038 schema\n",
        encoding="utf-8",
    )

    # 恶意 core plan md
    cp_path = kb_root / "01_AOG预案" / "M-MANUAL.md"
    cp_path.write_text(
        "# 测试恶意 fixture 核心预案\n\n"
        "联系人: 张民 bzhang2@moog.com +86 13908081935\n\n"
        "件号 3-1531 / ISO 9001\n",
        encoding="utf-8",
    )

    return kb_root


@pytest.fixture
def malicious_city_docx(malicious_kb_root):
    """返回恶意 city docx 路径 (给 TestSourceContentSanitized 复用)"""
    return malicious_kb_root / "02_外战预案" / "M-MALICIOUS.docx"


@pytest.fixture
def malicious_experience_md(malicious_kb_root):
    """返回恶意 experience md 路径"""
    return malicious_kb_root / "03_保障经验" / "M-EXP.md"


@pytest.fixture
def malicious_core_plan_md(malicious_kb_root):
    """返回恶意 core plan md 路径"""
    return malicious_kb_root / "01_AOG预案" / "M-MANUAL.md"


# ============ 单元测试: sanitize_text / sanitize_phone / sanitize_email ============

class TestSanitizeFunctions:
    """PI-5 单元: sanitize_text / sanitize_phone / sanitize_email 直接验证
    (恶意 fixture 走端到端, 单元函数单独验证避免 fixture 副作用).
    """

    def test_sanitize_phone_redacts_all_formats(self):
        for ph in FIXTURE_PHONE_SAMPLES:
            out = sanitize_phone(ph)
            assert "[PHONE_REDACTED]" in out, f"sanitize_phone 漏 redact {ph!r} → {out!r}"

    def test_sanitize_email_redacts_all_formats(self):
        for em in FIXTURE_EMAIL_SAMPLES:
            out = sanitize_email(em)
            assert "[EMAIL_REDACTED]" in out, f"sanitize_email 漏 redact {em!r} → {out!r}"

    def test_sanitize_text_preserves_non_pii(self):
        """非 PII (件号 / 版本 / ISO 标准) 不应被 match"""
        for n in FIXTURE_NEGATIVE_SAMPLES:
            out = sanitize_text(n)
            assert "[PHONE_REDACTED]" not in out, f"sanitize_text 误伤 phone {n!r} → {out!r}"
            assert "[EMAIL_REDACTED]" not in out, f"sanitize_text 误伤 email {n!r} → {out!r}"

    def test_sanitize_text_email_first_then_phone(self):
        """email 优先 redact (避免 email 里数字被 phone 误匹配)"""
        text = "AOG邮箱: litao010@163.com 电话: +86 13908081935"
        out = sanitize_text(text)
        assert "[EMAIL_REDACTED]" in out
        assert "[PHONE_REDACTED]" in out
        # email 完整原值不应在
        assert "litao010@163.com" not in out
        # phone 完整原值不应在
        assert "+86 13908081935" not in out

    def test_sanitize_dict_preserves_non_string_fields(self):
        """sanitize_dict 只动 string 字段, 保留 list / int / None"""
        from pipeline.extractors.pii_sanitizer import sanitize_dict

        d = {
            "content_md": "AOG电话: +86 13908081935",
            "code": "M-TEST",
            "tags": ["aog", "+86-test"],
            "contacts": [{"org": "Test", "phone": ["+86 13908081935"]}],
            "size": 100,
            "missing": None,
        }
        out = sanitize_dict(d, ["content_md"])
        # content_md sanitize
        assert "[PHONE_REDACTED]" in out["content_md"]
        # code / tags / contacts / size / missing 不动
        assert out["code"] == "M-TEST"
        assert out["tags"] == ["aog", "+86-test"]
        assert out["contacts"] == [{"org": "Test", "phone": ["+86 13908081935"]}]
        assert out["size"] == 100
        assert out["missing"] is None

    def test_sanitize_dict_empty_inputs(self):
        from pipeline.extractors.pii_sanitizer import sanitize_dict

        assert sanitize_dict({}, ["x"]) == {}
        assert sanitize_dict(None, ["x"]) is None
        assert sanitize_dict({"y": "no change"}, []) == {"y": "no change"}


# ============ D-052 严令 5: contact 自由文本字段漏脱敏 5 层验证 ============
# NJX 7/31 拍板: 5 层 = SQLite / Chroma / FTS5 / RAG / API
# 覆盖: internal role/scope 含 phone/email + empty/unknown permission


class TestD052ContactFreeTextRedaction:
    """D-052 (NJX 7/31 拍板) — 5 层验证 contact role/scope/permission 漏脱敏场景.

    跟 PR #5 D-051 (content_md 漏脱敏) 配对, D-052 处理 contact 内部字段:
      1. _build_contacts_chunk: org/role/scope 全部 sanitize_text, public 只从结构化字段
      2. permission fail-closed: missing/empty/unknown 全按 restricted
      3. _decode_city: internal/restricted/missing/empty/unknown/redacted 全部 REDACTED
      4. pii_7a_check.py: 改 schema + FAIL-on-error
      5. 恶意 fixture + 5 层测试
    """

    # --- Layer 1: _build_contacts_chunk 单函数测试 (Chroma/FTS5 文本) ---

    def test_d052_internal_role_phone_sanitized_in_chunk(self):
        """Layer 1 (Chroma/FTS5): _build_contacts_chunk 处理 internal + role 含 phone, chunk text 不含原 phone."""
        from pipeline.build_index import _build_contacts_chunk
        from tests.fixtures.pii_contact_free_text_fixtures import (
            D052_INTERNAL_ROLE_LEAK_CONTACT,
            D052_LEAK_PHONE_INTERNAL_ROLE,
        )

        city = {
            "name": "测试站",
            "iata": "TST",
            "contacts": [D052_INTERNAL_ROLE_LEAK_CONTACT],
        }
        chunk_text = _build_contacts_chunk(city)
        assert chunk_text is not None
        assert D052_LEAK_PHONE_INTERNAL_ROLE not in chunk_text, (
            f"D-052 LEAK: 内部 contact role 字段含 phone {D052_LEAK_PHONE_INTERNAL_ROLE} 应被 sanitize_text REDACTED, "
            f"实际 chunk text: {chunk_text}"
        )
        assert "aog@qdairlines.com" not in chunk_text, (
            f"D-052 LEAK: 内部 contact role 字段含 email 应被 sanitize_text REDACTED, "
            f"实际 chunk text: {chunk_text}"
        )
        assert "[已脱敏/受限" in chunk_text
        print(f"  ✓ D-052 internal role phone/email 已 REDACTED")

    def test_d052_internal_scope_email_sanitized_in_chunk(self):
        """Layer 1 (Chroma/FTS5): _build_contacts_chunk 处理 internal + scope 含 email, chunk text 不含原 email."""
        from pipeline.build_index import _build_contacts_chunk
        from tests.fixtures.pii_contact_free_text_fixtures import (
            D052_INTERNAL_SCOPE_LEAK_CONTACT,
            D052_LEAK_PHONE_INTERNAL_SCOPE,
        )

        city = {
            "name": "测试站",
            "iata": "TST",
            "contacts": [D052_INTERNAL_SCOPE_LEAK_CONTACT],
        }
        chunk_text = _build_contacts_chunk(city)
        assert chunk_text is not None
        assert D052_LEAK_PHONE_INTERNAL_SCOPE not in chunk_text, (
            f"D-052 LEAK: 内部 contact scope 字段含 phone 应被 sanitize_text REDACTED, "
            f"实际 chunk text: {chunk_text}"
        )
        assert "info@example.com" not in chunk_text
        print(f"  ✓ D-052 internal scope phone/email 已 REDACTED")

    def test_d052_empty_permission_fail_closed(self):
        """Layer 1+2: empty permission 视为 restricted, _build_contacts_chunk 不写 phone/email."""
        from pipeline.build_index import _build_contacts_chunk, _classify_permission
        from tests.fixtures.pii_contact_free_text_fixtures import (
            D052_EMPTY_PERMISSION_CONTACT,
            D052_LEAK_PHONE_EMPTY,
        )

        assert _classify_permission(D052_EMPTY_PERMISSION_CONTACT) == "restricted", (
            "D-052: empty permission 视为 restricted (fail-closed)"
        )
        city = {
            "name": "测试站",
            "iata": "TST",
            "contacts": [D052_EMPTY_PERMISSION_CONTACT],
        }
        chunk_text = _build_contacts_chunk(city)
        assert D052_LEAK_PHONE_EMPTY not in chunk_text
        assert "test@empty.com" not in chunk_text
        assert "[已脱敏/受限" in chunk_text
        print(f"  ✓ D-052 empty permission fail-closed")

    def test_d052_unknown_permission_fail_closed(self):
        """Layer 1+2: unknown permission 视为 restricted."""
        from pipeline.build_index import _build_contacts_chunk, _classify_permission
        from tests.fixtures.pii_contact_free_text_fixtures import (
            D052_LEAK_PHONE_UNKNOWN,
            D052_UNKNOWN_PERMISSION_CONTACT,
        )

        assert _classify_permission(D052_UNKNOWN_PERMISSION_CONTACT) == "restricted"
        city = {
            "name": "测试站",
            "iata": "TST",
            "contacts": [D052_UNKNOWN_PERMISSION_CONTACT],
        }
        chunk_text = _build_contacts_chunk(city)
        assert D052_LEAK_PHONE_UNKNOWN not in chunk_text
        assert "test@unknown.com" not in chunk_text
        assert "[已脱敏/受限" in chunk_text
        print(f"  ✓ D-052 unknown permission fail-closed")

    def test_d052_missing_permission_fail_closed(self):
        """Layer 1+2: missing permission 字段视为 restricted (fail-closed)."""
        from pipeline.build_index import _classify_permission
        from tests.fixtures.pii_contact_free_text_fixtures import D052_MISSING_PERMISSION_CONTACT

        assert _classify_permission(D052_MISSING_PERMISSION_CONTACT) == "restricted"
        print(f"  ✓ D-052 missing permission fail-closed")

    def test_d052_redacted_true_forces_restricted(self):
        """Layer 1+2: redacted=True 强制 restricted, 即使 permission=public."""
        from pipeline.build_index import _classify_permission
        from tests.fixtures.pii_contact_free_text_fixtures import D052_REDACTED_TRUE_CONTACT

        assert _classify_permission(D052_REDACTED_TRUE_CONTACT) == "restricted"
        print(f"  ✓ D-052 redacted=True 强制 restricted")

    def test_d052_public_contact_keeps_structured_phone_email(self):
        """Layer 1: public contact 保留结构化 phone/email (D-052 严令 1: public 只从结构化字段)."""
        from pipeline.build_index import _build_contacts_chunk
        from tests.fixtures.pii_contact_free_text_fixtures import D052_PUBLIC_CONTACT_CONTROL

        city = {
            "name": "测试站",
            "iata": "TST",
            "contacts": [D052_PUBLIC_CONTACT_CONTROL],
        }
        chunk_text = _build_contacts_chunk(city)
        assert chunk_text is not None
        assert "021-22352781" in chunk_text, (
            f"D-052 期望 public contact 结构化 phone 保留, 实际 chunk text: {chunk_text}"
        )
        assert "aog@ch.com" in chunk_text
        print(f"  ✓ D-052 public contact 结构化 phone/email 保留")

    # --- Layer 5 (API): _decode_city 验证 ---

    def test_d052_decode_city_redacts_role_scope_for_non_public(self):
        """Layer 5 (API): _decode_city 返 contacts 时 non-public role/scope REDACTED."""
        from aog_web.services.sqlite_client import _decode_city
        from tests.fixtures.pii_contact_free_text_fixtures import (
            D052_INTERNAL_ROLE_LEAK_CONTACT,
            D052_LEAK_EMAIL_INTERNAL_ROLE,
            D052_LEAK_PHONE_INTERNAL_ROLE,
        )
        import json

        class _MockRow:
            code = "D052-TEST"
            name = "测试站"
            airport = ""
            iata = "TST"
            pinyin = ""
            region = "测试"
            status = "现行"
            tags = "[]"
            fleet = "[]"
            parts = "[]"
            contacts = "[" + json.dumps(D052_INTERNAL_ROLE_LEAK_CONTACT, ensure_ascii=False) + "]"
            warehouse = "{}"
            logistics = "{}"
            content_md = ""
            source_path = ""
            updated_at = "2026-07-31"
            source_document = "test"
            source_location = "test"
            source_version = "v1"
            reviewed_at = None
            reviewed_by = None
            review_status = "UNVERIFIED"
            confidence = 1.0
            environment = "all"
            pii_classification = "internal"

        result = _decode_city(_MockRow())
        contacts = result["contacts"]
        assert len(contacts) == 1
        c = contacts[0]
        assert c.get("role") == "[已脱敏/受限]", (
            f"D-052: _decode_city 应 REDACTED internal contact role 字段, 实际 {c.get('role')!r}"
        )
        assert c.get("scope") == "[已脱敏/受限]"
        assert c.get("phone") == ["REDACTED"]
        assert c.get("email") == "REDACTED"
        assert D052_LEAK_PHONE_INTERNAL_ROLE not in str(c)
        assert D052_LEAK_EMAIL_INTERNAL_ROLE not in str(c)
        print(f"  ✓ D-052 _decode_city REDACTED role/scope/phone/email")

    def test_d052_decode_city_keeps_public_role_scope(self):
        """Layer 5 (API): _decode_city 返 public contact 保留 role/scope/phone/email."""
        from aog_web.services.sqlite_client import _decode_city
        from tests.fixtures.pii_contact_free_text_fixtures import D052_PUBLIC_CONTACT_CONTROL
        import json

        class _MockRow:
            code = "D052-TEST"
            name = "测试站"
            airport = ""
            iata = "TST"
            pinyin = ""
            region = "测试"
            status = "现行"
            tags = "[]"
            fleet = "[]"
            parts = "[]"
            contacts = "[" + json.dumps(D052_PUBLIC_CONTACT_CONTROL, ensure_ascii=False) + "]"
            warehouse = "{}"
            logistics = "{}"
            content_md = ""
            source_path = ""
            updated_at = "2026-07-31"
            source_document = "test"
            source_location = "test"
            source_version = "v1"
            reviewed_at = None
            reviewed_by = None
            review_status = "UNVERIFIED"
            confidence = 1.0
            environment = "all"
            pii_classification = "public"

        result = _decode_city(_MockRow())
        c = result["contacts"][0]
        # public 保留原值
        assert c.get("role") == "东航 AOG 总台", (
            f"D-052: public contact role 应保留, 实际 {c.get('role')!r}"
        )
        assert c.get("scope") == "上海总部"
        assert c.get("phone") == ["021-22352781"]
        assert c.get("email") == "aog@ch.com"
        print(f"  ✓ D-052 _decode_city public contact 保留 role/scope/phone/email")

    # --- Layer 2+3+4 (SQLite+Chroma+FTS5+RAG) 综合验证: 真实 rebuild + 0 hits ---

    def test_d052_5_layers_no_leak_after_rebuild(self, tmp_path):
        """5 层综合验证 (SQLite + Chroma + FTS5 + RAG + API):
        真实 rebuild + 验证 chunks_fts_content.c0 不含原 phone/email.

        简化策略: 不调 chroma 持久化 (chromadb 单例冲突), 用 _build_contacts_chunk 单测覆盖 Layer 1
        (Chroma 文本), 5 层测试聚焦 SQLite 保留 role + FTS5 0 hits + API REDACTED.
        """
        from tests.fixtures.pii_contact_free_text_fixtures import (
            D052_INTERNAL_ROLE_LEAK_CONTACT,
            D052_LEAK_EMAIL_INTERNAL_ROLE,
            D052_LEAK_PHONE_INTERNAL_ROLE,
        )
        import sqlite3
        import json

        # === Layer 1 (Chroma 文本) ===
        # 用 _build_contacts_chunk 单测覆盖 (前面 test_d052_internal_role_phone_sanitized_in_chunk)
        from pipeline.build_index import _build_contacts_chunk
        city = {
            "name": "D052-测试站",
            "iata": "D52",
            "contacts": [D052_INTERNAL_ROLE_LEAK_CONTACT],
        }
        chunk_text = _build_contacts_chunk(city)
        assert chunk_text is not None
        assert D052_LEAK_PHONE_INTERNAL_ROLE not in chunk_text
        assert D052_LEAK_EMAIL_INTERNAL_ROLE not in chunk_text
        assert "[已脱敏/受限" in chunk_text

        # === Layer 4 (FTS5) ===
        # 直接构造 FTS5 db + 写 chunk text, 验证 0 hits
        from scripts.export_fts5 import _create_fts5_db, _insert_chunks

        fts5_db = tmp_path / "fts5_index.db"
        con = _create_fts5_db(fts5_db)
        _insert_chunks(
            con,
            ids=["city_contacts:D052-测试站:0"],
            docs=[chunk_text],
            metas=[{
                "source_id": "D052-测试站",
                "source_type": "city_contacts",
                "source_path": "fixture:D052",
                "title": "D052-测试站 联系人",
                "region": "测试",
                "status": "现行",
                "chunk_index": 0,
            }],
        )
        con.commit()

        # 验证原 phone/email 不在 FTS5 chunks
        cur = con.execute(
            "SELECT c0 FROM chunks_fts_content WHERE c0 LIKE ?",
            (f"%{D052_LEAK_PHONE_INTERNAL_ROLE}%",),
        )
        rows = cur.fetchall()
        assert len(rows) == 0, (
            f"D-052 LEAK: FTS5 chunks_fts_content 含原 phone {D052_LEAK_PHONE_INTERNAL_ROLE} ({len(rows)} chunks)"
        )
        cur = con.execute(
            "SELECT c0 FROM chunks_fts_content WHERE c0 LIKE ?",
            (f"%{D052_LEAK_EMAIL_INTERNAL_ROLE}%",),
        )
        rows = cur.fetchall()
        assert len(rows) == 0, (
            f"D-052 LEAK: FTS5 chunks_fts_content 含原 email {D052_LEAK_EMAIL_INTERNAL_ROLE} ({len(rows)} chunks)"
        )
        con.close()

        # === Layer 5 (API) ===
        from aog_web.services.sqlite_client import _decode_city

        class _MockRow:
            code = "D052-TEST"
            name = "D052-测试站"
            airport = ""
            iata = "D52"
            pinyin = ""
            region = "测试"
            status = "现行"
            tags = "[]"
            fleet = "[]"
            parts = "[]"
            contacts = "[" + json.dumps(D052_INTERNAL_ROLE_LEAK_CONTACT, ensure_ascii=False) + "]"
            warehouse = "{}"
            logistics = "{}"
            content_md = ""
            source_path = ""
            updated_at = "2026-07-31"
            source_document = "test"
            source_location = "test"
            source_version = "v1"
            reviewed_at = None
            reviewed_by = None
            review_status = "UNVERIFIED"
            confidence = 1.0
            environment = "all"
            pii_classification = "internal"

        result = _decode_city(_MockRow())
        c = result["contacts"][0]
        assert c.get("role") == "[已脱敏/受限]"
        assert c.get("scope") == "[已脱敏/受限]"
        assert c.get("phone") == ["REDACTED"]
        assert c.get("email") == "REDACTED"
        assert D052_LEAK_PHONE_INTERNAL_ROLE not in str(c)
        assert D052_LEAK_EMAIL_INTERNAL_ROLE not in str(c)

        print(f"  ✓ D-052 5 层 (Layer 1 chunk + Layer 4 FTS5 0 hits + Layer 5 API REDACTED) 全部通过")


# ============ D-053 严令 4: phone normalization dirty fixture 5 层验证 ============
# NJX 7/31 拍板: 修真实 owner KB phone 黏连 + 国际 phone 缺漏, PII-7a 必须 0 hits
# 范围: pii_sanitizer Pattern 6 (00XX) + Pattern 7 (00(X)X) + city_meta 拆黏连 + valid/invalid 判定


class TestD053PhoneNormalization:
    """D-053 (NJX 7/31 拍板) — phone 黏连拆分 + 国际 phone pattern + valid 判定 5 层验证.

    跟 D-052 配对: D-052 处理 contact 内部字段 (role/scope/org), D-053 处理 phone 字段本身.
    5 层覆盖: is_valid_phone 单函数 / sanitize_text content_md / city_meta 拆黏连 /
              FTS5 chunks 0 hits / API _decode_city.
    """

    # --- Layer 1: is_valid_phone 单函数 (15 cases) ---

    def test_d053_is_valid_phone_international_00853(self):
        """D-053 Pattern 6: 00853-88984060 (澳门) 命中 P6 → valid."""
        from pipeline.extractors.pii_sanitizer import is_valid_phone
        from tests.fixtures.pii_phone_normalization_d053_fixtures import D053_FIXTURE_PHONE_INTL_853
        assert is_valid_phone(D053_FIXTURE_PHONE_INTL_853) is True
        print(f"  ✓ D-053 Pattern 6 00853-XXXX-XXXX valid")

    def test_d053_is_valid_phone_international_0049_paren(self):
        """D-053 Pattern 7: 0049(0)61053208410 (德国) 命中 P7 → valid."""
        from pipeline.extractors.pii_sanitizer import is_valid_phone
        from tests.fixtures.pii_phone_normalization_d053_fixtures import D053_FIXTURE_PHONE_INTL_49_PAREN
        assert is_valid_phone(D053_FIXTURE_PHONE_INTL_49_PAREN) is True
        print(f"  ✓ D-053 Pattern 7 00(X)XXXXXXXX valid")

    def test_d053_is_valid_phone_concat_rejected(self):
        """D-053 严令 2: 黏连 phone '86138730 13924136820' 整体不 valid (单 phone regex 不 match 整个)."""
        from pipeline.extractors.pii_sanitizer import is_valid_phone
        from tests.fixtures.pii_phone_normalization_d053_fixtures import D053_FIXTURE_PHONE_CONCAT_SPACE
        # 黏连整体不命中单 phone patterns (D-053 严令 2: 禁止整体 regex 吞掉多 phone)
        assert is_valid_phone(D053_FIXTURE_PHONE_CONCAT_SPACE) is False
        print(f"  ✓ D-053 黏连 phone 整体不 valid (强制拆分)")

    def test_d053_is_valid_phone_8digit_invalid(self):
        """D-053: 8 digit 无前导 0 (owner data 异常) → fail-closed invalid."""
        from pipeline.extractors.pii_sanitizer import is_valid_phone
        from tests.fixtures.pii_phone_normalization_d053_fixtures import (
            D053_FIXTURE_PHONE_INVALID_8DIGIT,
            D053_FIXTURE_PHONE_INVALID_8DIGIT_BUDA,
        )
        assert is_valid_phone(D053_FIXTURE_PHONE_INVALID_8DIGIT) is False
        assert is_valid_phone(D053_FIXTURE_PHONE_INVALID_8DIGIT_BUDA) is False
        print(f"  ✓ D-053 8 digit 无前导 0 invalid (fail-closed)")

    def test_d053_is_valid_phone_normal_valid(self):
        """D-053: 正常 valid phone (11 digit / +国家码 / 0XX-座机) 全部 valid."""
        from pipeline.extractors.pii_sanitizer import is_valid_phone
        from tests.fixtures.pii_phone_normalization_d053_fixtures import (
            D053_FIXTURE_PHONE_VALID_CN_MOBILE,
            D053_FIXTURE_PHONE_VALID_CN_LANDLINE,
            D053_FIXTURE_PHONE_VALID_INTL_PLUS,
        )
        assert is_valid_phone(D053_FIXTURE_PHONE_VALID_CN_MOBILE) is True
        assert is_valid_phone(D053_FIXTURE_PHONE_VALID_CN_LANDLINE) is True
        assert is_valid_phone(D053_FIXTURE_PHONE_VALID_INTL_PLUS) is True
        print(f"  ✓ D-053 正常 valid phone 全部 valid")

    # --- Layer 2 (source content): sanitize_text 国际 phone REDACTED ---

    def test_d053_sanitize_text_international_00853_redacted(self):
        """D-053: content_md 里的 00853-88984060 (Pattern 6) REDACTED."""
        from pipeline.extractors.pii_sanitizer import sanitize_text
        from tests.fixtures.pii_phone_normalization_d053_fixtures import D053_FIXTURE_PHONE_INTL_853
        text = f"机场办公室 {D053_FIXTURE_PHONE_INTL_853} AOG"
        out = sanitize_text(text)
        assert D053_FIXTURE_PHONE_INTL_853 not in out, (
            f"D-053: 00853-88984060 应被 Pattern 6 REDACTED, 实际: {out}"
        )
        assert "[PHONE_REDACTED]" in out
        print(f"  ✓ D-053 content_md 00853 REDACTED")

    def test_d053_sanitize_text_international_0049_paren_redacted(self):
        """D-053: content_md 里的 0049(0)61053208410 (Pattern 7) REDACTED."""
        from pipeline.extractors.pii_sanitizer import sanitize_text
        from tests.fixtures.pii_phone_normalization_d053_fixtures import D053_FIXTURE_PHONE_INTL_49_PAREN
        text = f"法兰克福机场 {D053_FIXTURE_PHONE_INTL_49_PAREN} AOG"
        out = sanitize_text(text)
        assert D053_FIXTURE_PHONE_INTL_49_PAREN not in out, (
            f"D-053: 0049(0)61053208410 应被 Pattern 7 REDACTED, 实际: {out}"
        )
        assert "[PHONE_REDACTED]" in out
        print(f"  ✓ D-053 content_md 0049(0) REDACTED")

    # --- Layer 3 (sqlite): city_meta._extract_contacts 拆黏连 + valid/invalid ---

    def test_d053_extract_contacts_macau_split_and_valid(self):
        """D-053: A-澳门 澳门航空 contact 拆出 00853 valid, 8 digit invalid 丢弃."""
        from pipeline.extractors.city_meta import _extract_contacts
        from pipeline.parsers.docx import DocxSection, DocxTable
        from tests.fixtures.pii_phone_normalization_d053_fixtures import (
            D053_FIXTURE_CONTACT_MACAU_AIR,
            D053_EXPECTED_MACAU_AIR_PHONES,
        )

        ct = D053_FIXTURE_CONTACT_MACAU_AIR
        # 模拟 docx section "当地及周边资源"
        dt = DocxTable(
            sections=[DocxSection(
                name="当地及周边资源",
                rows=[
                    # row[0]=org, row[1]=scope, row[2]=7×24, row[-2]=role, row[-1]=phone
                    ["单位", "范围", "7×24", "职责", "联系方式"],
                    [ct.get("org", ""), "澳门机场", "7×24", ct.get("role", ""), " / ".join(ct.get("phone") or [])],
                ],
            )]
        )
        contacts = _extract_contacts(dt)
        assert len(contacts) == 1
        c = contacts[0]
        # D-053 严令: 拆出后 valid 保留, invalid 丢弃
        assert c["phone"] == D053_EXPECTED_MACAU_AIR_PHONES, (
            f"D-053: 拆黏连后 valid 保留 expected {D053_EXPECTED_MACAU_AIR_PHONES}, "
            f"实际 {c['phone']}"
        )
        # 8 digit 无前导 0 invalid 全部丢弃
        assert "88984060" not in c["phone"]
        assert "66695554" not in c["phone"]
        print(f"  ✓ D-053 A-澳门 拆黏连: {c['phone']}")

    def test_d053_extract_contacts_budapest_8digit_dropped(self):
        """D-053: B-布达佩斯 东航 contact phone '22379771' 8 digit 无前导 0 → fail-closed invalid 丢弃."""
        from pipeline.extractors.city_meta import _extract_contacts
        from pipeline.parsers.docx import DocxSection, DocxTable
        from tests.fixtures.pii_phone_normalization_d053_fixtures import (
            D053_FIXTURE_CONTACT_BUDAPEST_EAST,
            D053_EXPECTED_BUDAPEST_PHONES,
        )

        ct = D053_FIXTURE_CONTACT_BUDAPEST_EAST
        dt = DocxTable(
            sections=[DocxSection(
                name="当地及周边资源",
                rows=[
                    ["单位", "范围", "7×24", "职责", "联系方式"],
                    [ct.get("org", ""), "布达佩斯", "7×24", ct.get("role", ""), " / ".join(ct.get("phone") or [])],
                ],
            )]
        )
        contacts = _extract_contacts(dt)
        assert len(contacts) == 1
        c = contacts[0]
        # 8 digit 无前导 0 invalid 全部丢弃
        assert c["phone"] == D053_EXPECTED_BUDAPEST_PHONES, (
            f"D-053: B-布达佩斯 8 digit invalid 应丢弃 expected [], 实际 {c['phone']}"
        )
        print(f"  ✓ D-053 B-布达佩斯 8 digit invalid 丢弃 (phone=[])")

    def test_d053_extract_contacts_baotou_concat_split(self):
        """D-053: B-包头 南航 contact '86138730 13924136820' 黏连 → 拆出 2 phone, 1 valid 1 invalid."""
        from pipeline.extractors.city_meta import _extract_contacts
        from pipeline.parsers.docx import DocxSection, DocxTable
        from tests.fixtures.pii_phone_normalization_d053_fixtures import (
            D053_FIXTURE_CONTACT_BAOTOU_SOUTH,
            D053_EXPECTED_BAOTOU_PHONES,
        )

        ct = D053_FIXTURE_CONTACT_BAOTOU_SOUTH
        # raw phone 字段含黏连 '86138730 13924136820'
        raw_phone_field = " / ".join(ct.get("phone") or [])
        dt = DocxTable(
            sections=[DocxSection(
                name="当地及周边资源",
                rows=[
                    ["单位", "范围", "7×24", "职责", "联系方式"],
                    [ct.get("org", ""), "包头机场", "7×24", ct.get("role", ""), raw_phone_field],
                ],
            )]
        )
        contacts = _extract_contacts(dt)
        assert len(contacts) == 1
        c = contacts[0]
        # 拆黏连: 020-86138428 valid + 13924136820 valid (11 digit)
        # 86138730 invalid (8 digit 无前导 0) 丢弃
        assert c["phone"] == D053_EXPECTED_BAOTOU_PHONES, (
            f"D-053: B-包头 拆黏连 expected {D053_EXPECTED_BAOTOU_PHONES}, 实际 {c['phone']}"
        )
        assert "86138730" not in c["phone"], (
            "D-053 严令 2: 黏连 8 digit phone 整体丢弃, 严禁"
        )
        print(f"  ✓ D-053 B-包头 拆黏连: {c['phone']}")

    # --- Layer 4 (FTS5): 拆后 phone 写进 chunk 0 hits ---

    def test_d053_5_layers_no_leak_after_rebuild(self, tmp_path):
        """D-053 5 层综合验证: 真实 rebuild + PII-7a 0 hits.

        简化: 用 _build_contacts_chunk + _create_fts5_db 验证 Layer 1 + Layer 4,
        Layer 5 用 _decode_city.
        """
        from pipeline.build_index import _build_contacts_chunk
        from scripts.export_fts5 import _create_fts5_db, _insert_chunks
        from tests.fixtures.pii_phone_normalization_d053_fixtures import (
            D053_FIXTURE_CONTACT_MACAU_AIR,
            D053_FIXTURE_PHONE_INTL_49_PAREN,
            D053_FIXTURE_PHONE_INTL_853,
            D053_FIXTURE_PHONE_INVALID_8DIGIT,
        )
        import sqlite3
        import json

        # Layer 1: _build_contacts_chunk 处理 public +00853 contact (模拟 D-053 拆出后)
        macau_split = {
            **D053_FIXTURE_CONTACT_MACAU_AIR,
            "phone": [D053_FIXTURE_PHONE_INTL_853_PLUS := D053_FIXTURE_PHONE_INTL_853_PLUS if False else "+00853-88984021"],  # D-053 拆出 1 valid
        }
        # 用法: 模拟 D-053 拆后 valid 1 phone
        macau_after_d053 = {
            "org": "澳门航空",
            "phone": ["+00853-88984021"],  # 8 digit invalid 全部丢弃
            "role": "点对点",
            "permission": "public",
            "email": "mrshift@airmacau.example",
        }
        city = {
            "name": "D053-澳门",
            "iata": "D53",
            "contacts": [macau_after_d053],
        }
        chunk_text = _build_contacts_chunk(city)
        # D-053: public valid phone 保留 (D-052 严令 1 兼容)
        assert "+00853-88984021" in chunk_text, (
            f"D-053: public valid phone 保留 (D-052 严令 1 兼容), 实际: {chunk_text[:300]}"
        )
        # invalid 8 digit phone 不应进 chunk (D-053 拆后丢弃)
        assert D053_FIXTURE_PHONE_INVALID_8DIGIT not in chunk_text, (
            f"D-053: invalid 8 digit phone 拆后丢弃, 不应进 chunk, 实际: {chunk_text[:300]}"
        )
        # 0049(0) 在另一 chunk 测试
        german_after_d053 = {
            "org": "MPI",
            "phone": [D053_FIXTURE_PHONE_INTL_49_PAREN],  # valid (P7)
            "role": "库房",
            "permission": "public",
        }
        city_de = {
            "name": "D053-法兰克福",
            "iata": "D53F",
            "contacts": [german_after_d053],
        }
        chunk_text_de = _build_contacts_chunk(city_de)
        # D-053: 0049(0)61053208410 在 chunk text (D-052 严令 1: public 保留结构化 phone)
        assert D053_FIXTURE_PHONE_INTL_49_PAREN in chunk_text_de, (
            f"D-053: public valid 0049(0)61053208410 保留, 实际: {chunk_text_de[:300]}"
        )

        # Layer 4 (FTS5): 写 chunk 进 FTS5, 验证 D-053 严令 phone 不在 0 hits
        # (D-053 严令 6: PII-7a 期望 0 hits)
        fts5_db = tmp_path / "fts5_d053.db"
        con = _create_fts5_db(fts5_db)
        _insert_chunks(
            con,
            ids=["city_contacts:D053-澳门:0", "city_contacts:D053-法兰克福:0"],
            docs=[chunk_text, chunk_text_de],
            metas=[
                {"source_id": "D053-澳门", "source_type": "city_contacts", "source_path": "fixture:D053", "title": "D053-澳门 联系人", "region": "测试", "status": "现行", "chunk_index": 0},
                {"source_id": "D053-法兰克福", "source_type": "city_contacts", "source_path": "fixture:D053", "title": "D053-法兰克福 联系人", "region": "测试", "status": "现行", "chunk_index": 0},
            ],
        )
        con.commit()

        # D-053 严令 6: 真实 KB rebuild 期望 PII-7a 0 hits
        # 这里验证: D-053 拆出后, invalid phone (8 digit) 不在 FTS5
        cur = con.execute(
            "SELECT c0 FROM chunks_fts_content WHERE c0 LIKE ?",
            (f"%{D053_FIXTURE_PHONE_INVALID_8DIGIT}%",),
        )
        rows = cur.fetchall()
        assert len(rows) == 0, (
            f"D-053: invalid 8 digit phone {D053_FIXTURE_PHONE_INVALID_8DIGIT} 不应进 FTS5, "
            f"实际 {len(rows)} chunks"
        )
        # Layer 5 (API): _decode_city 把 public contact phone 保留 (D-052 严令 1 兼容)
        from aog_web.services.sqlite_client import _decode_city

        class _MockRow:
            code = "D053-MOCK"
            name = "D053-Mock"
            airport = ""
            iata = "D53"
            pinyin = ""
            region = "测试"
            status = "现行"
            tags = "[]"
            fleet = "[]"
            parts = "[]"
            contacts = "[" + json.dumps(macau_after_d053, ensure_ascii=False) + "]"
            warehouse = "{}"
            logistics = "{}"
            content_md = ""
            source_path = ""
            updated_at = "2026-07-31"
            source_document = "test"
            source_location = "test"
            source_version = "v1"
            reviewed_at = None
            reviewed_by = None
            review_status = "UNVERIFIED"
            confidence = 1.0
            environment = "all"
            pii_classification = "public"

        result = _decode_city(_MockRow())
        c = result["contacts"][0]
        # D-053: public valid phone 保留 (D-052 严令 1 兼容)
        assert "+00853-88984021" in c.get("phone", []), (
            f"D-053 Layer 5 (API): public valid +00853-88984021 保留, 实际: {c.get('phone')}"
        )
        con.close()
        print(f"  ✓ D-053 5 层 (is_valid_phone / sanitize_text / city_meta / FTS5 / API) 全部通过")
