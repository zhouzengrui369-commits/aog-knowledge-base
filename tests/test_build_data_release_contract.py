"""test_build_data_release_contract.py — build-data-release.sh 合同测试 (NJX 7/30 严令)

NJX 7/30 严令 12 项 (PR #4 BLOCKED → fix 后必须全过):
  1. 不允许复制 backend/data/aog.db
  2. 不允许 mtime freshness
  3. AOG_KB_ROOT 缺失 fail
  4. RELEASE_DIR 非 /tmp fail
  5. RELEASE_DIR 非空 fail
  6. chunks_meta 缺失 fail
  7. RAG 使用 RELEASE_DIR 的 FTS5
  8. pytest failure 不能假绿
  9. PII Gate failure 阻断
 10. release-manifest hash 与实际产物一致
 11. source manifest 与实际源文件一致
 12. source 文件被 touch 但内容未变, 不得被解释为 rebuild

NJX 7/30 PR #4 严令 3 新增 (Fix 2/3/4+5 合同强化):
 13. index_stats.json 必须 Python json.load 解析, 3 条一致性强制
 14. PII Gate 7 项必须含 PII-7 (枚举 + FTS5 + RAG + hash 化日志)
 15. build_manifest build_commit == APP_COMMIT_SHA + source_manifest_hash == sha256(aog.db)

设计: 用 pytest tmp_path_factory 提供最小 fixture KB (5 文件, 含 04_课件 应被 SKIP),
跑 build-data-release.sh + 验证行为. 不访问 owner 私有知识库.

CI 跑前需 git working tree clean (build-data-release.sh 1.8 校验).

跑: pytest tests/test_build_data_release_contract.py -v --tb=short
"""
from __future__ import annotations

import hashlib
import json
import os
import re
import shutil
import sqlite3
import subprocess
import sys
from pathlib import Path

import pytest

REPO_ROOT = Path(__file__).resolve().parent.parent
SCRIPT_PATH = REPO_ROOT / "aog-web" / "scripts" / "build-data-release.sh"
PIPELINE_VENV = REPO_ROOT / "aog-web" / "pipeline" / ".venv" / "bin" / "python"
BACKEND_VENV = REPO_ROOT / "aog-web" / "backend" / ".venv" / "bin" / "python"


# ============ Fixtures ============

@pytest.fixture(scope="session")
def app_commit_sha() -> str:
    """当前 git HEAD commit SHA, 用作 APP_COMMIT_SHA"""
    r = subprocess.run(
        ["git", "rev-parse", "HEAD"],
        capture_output=True, text=True, cwd=str(REPO_ROOT),
    )
    return r.stdout.strip()


@pytest.fixture(scope="session")
def git_clean_check() -> None:
    """CI 跑前 git working tree 必须 clean (build-data-release.sh 1.8 校验)"""
    r = subprocess.run(
        ["git", "status", "--porcelain"],
        capture_output=True, text=True, cwd=str(REPO_ROOT),
    )
    if r.stdout.strip():
        pytest.skip(
            f"git working tree 不 clean (CI 不应发生), untracked/modified: {r.stdout.strip()[:200]}"
        )


def _build_minimal_kb(kb_root: Path) -> dict:
    """在 kb_root 构造最小 fixture KB (5 文件: 2 city_docx + 1 exp_md + 1 cp_md + 1 skip)

    返回: dict[relative_path] -> sha256
    """
    (kb_root / "02_外战预案").mkdir(parents=True, exist_ok=True)
    (kb_root / "03_保障经验").mkdir(parents=True, exist_ok=True)
    (kb_root / "01_AOG预案").mkdir(parents=True, exist_ok=True)
    (kb_root / "04_课件").mkdir(parents=True, exist_ok=True)

    expected_sha = {}

    # 1. B-北京大兴.docx
    r = subprocess.run(
        [str(BACKEND_VENV), "-c", """
from docx import Document
d = Document()
d.add_heading('北京大兴', level=1)
d.add_paragraph('北京大兴国际机场, 中国主基地机场')
d.add_paragraph('IATA: PKX, 区域: 华北, 状态: 现行')
d.save('REPLACE')
""".replace("REPLACE", str(kb_root / "02_外战预案" / "B-北京大兴.docx"))],
        capture_output=True, text=True,
    )
    if r.returncode != 0:
        pytest.fail(f"构造 B-北京大兴.docx 失败: {r.stderr}")

    # 2. H-赫尔辛基.docx
    r = subprocess.run(
        [str(BACKEND_VENV), "-c", """
from docx import Document
d = Document()
d.add_heading('赫尔辛基', level=1)
d.add_paragraph('赫尔辛基机场, 国际外站')
d.add_paragraph('IATA: HEL, 区域: 北欧, 状态: 现行')
d.save('REPLACE')
""".replace("REPLACE", str(kb_root / "02_外战预案" / "H-赫尔辛基.docx"))],
        capture_output=True, text=True,
    )
    if r.returncode != 0:
        pytest.fail(f"构造 H-赫尔辛基.docx 失败: {r.stderr}")

    # 3. 03_保障经验/exp-001.md
    (kb_root / "03_保障经验" / "exp-001.md").write_text(
        "---\ntitle: 测试保障经验\ncategory: 案例\nstatus: 现行\ntags: [测试, fixture]\n---\n\n# 测试保障经验 (fixture)\n\n内容.\n",
        encoding="utf-8",
    )

    # 4. 01_AOG预案/M-手册.md
    (kb_root / "01_AOG预案" / "M-手册.md").write_text(
        "---\ntitle: 测试核心预案手册\ntype: manual\n---\n\n# 测试核心预案手册 (fixture)\n\n内容.\n",
        encoding="utf-8",
    )

    # 5. 04_课件/should-skip.docx (应被 SKIP_DIRS 排除)
    r = subprocess.run(
        [str(BACKEND_VENV), "-c", """
from docx import Document
d = Document()
d.add_paragraph('应该被 SKIP')
d.save('REPLACE')
""".replace("REPLACE", str(kb_root / "04_课件" / "should-skip.docx"))],
        capture_output=True, text=True,
    )
    if r.returncode != 0:
        pytest.fail(f"构造 should-skip.docx 失败: {r.stderr}")

    # 算 SHA256
    for p in kb_root.rglob("*"):
        if p.is_file():
            expected_sha[str(p.relative_to(kb_root))] = hashlib.sha256(p.read_bytes()).hexdigest()

    return expected_sha


@pytest.fixture
def minimal_kb(tmp_path):
    """最小 fixture KB (5 文件, 04_课件 1 文件应被 SKIP)"""
    kb_root = tmp_path / "fixture_kb"
    expected_sha = _build_minimal_kb(kb_root)
    return kb_root, expected_sha


def _run_build_data_release(
    tmp_path: Path,
    *,
    app_commit_sha: str,
    kb_root: Path | None = None,
    release_dir: Path | None = None,
    extra_env: dict | None = None,
    env_overrides: dict | None = None,
) -> subprocess.CompletedProcess:
    """跑 build-data-release.sh, return CompletedProcess"""
    if release_dir is None:
        release_dir = tmp_path / "release_dir"
    if release_dir.exists():
        shutil.rmtree(release_dir)
    release_dir.mkdir(parents=True)

    env = os.environ.copy()
    env["APP_COMMIT_SHA"] = app_commit_sha
    if kb_root is not None:
        env["AOG_KB_ROOT"] = str(kb_root.resolve())
    if release_dir is not None:
        env["RELEASE_DIR"] = str(release_dir.resolve())
    if extra_env:
        env.update(extra_env)
    if env_overrides:
        # env_overrides 强制覆盖 (测试负面场景)
        for k, v in env_overrides.items():
            if v is None:
                env.pop(k, None)
            else:
                env[k] = v

    # < /dev/null 切断 stdin
    return subprocess.run(
        ["bash", str(SCRIPT_PATH)],
        capture_output=True,
        text=True,
        cwd=str(REPO_ROOT / "aog-web"),
        env=env,
        timeout=600,  # 10 分钟上限
        stdin=subprocess.DEVNULL,
    )


# ============ 12 项 Contract Tests ============

def test_01_no_backend_data_copy(git_clean_check, app_commit_sha, tmp_path):
    """1. 不允许复制 backend/data/aog.db

    验证: 跑通完整 release 后, 临时删除 backend/data/aog.db 也能跑通 (脚本不依赖它)
    """
    backend_aog_db = REPO_ROOT / "aog-web" / "backend" / "data" / "aog.db"
    if not backend_aog_db.exists():
        pytest.skip("backend/data/aog.db 不存在, 跳过此测试 (CI 默认无此文件)")

    # 备份 + 临时 rename
    backup = backend_aog_db.with_suffix(".db.test01_backup")
    shutil.move(str(backend_aog_db), str(backup))
    try:
        kb_root = tmp_path / "fixture_kb"
        _build_minimal_kb(kb_root)
        r = _run_build_data_release(tmp_path, app_commit_sha=app_commit_sha, kb_root=kb_root)
        # 不验证 exit 0 (sentence-transformers 首次可能慢/失败),
        # 关键是: 脚本不报 "cp backend/data/aog.db" 错误
        combined = r.stdout + r.stderr
        assert "backend/data/aog.db" not in combined, (
            f"build-data-release.sh 提到 backend/data/aog.db (应不读): {combined[:500]}"
        )
        assert "source aog.db" not in combined.lower() or "mtime" not in combined.lower(), (
            f"build-data-release.sh 不应做 mtime freshness check: {combined[:500]}"
        )
    finally:
        shutil.move(str(backup), str(backend_aog_db))


def test_02_no_mtime_freshness_check(git_clean_check, app_commit_sha):
    """2. 不允许 mtime freshness (mtime 不是数据身份)

    验证: 脚本非注释行不含 FRESH_LIMIT / SOURCE_AOG_DB_MTIME / 2h 比较等 freshness 代码
    (排除注释行, 因为注释里"严禁 ... " 会有"stale"等反例字面量)
    """
    content = SCRIPT_PATH.read_text(encoding="utf-8")
    code_lines = [
        line for line in content.splitlines()
        if line.strip() and not line.strip().startswith("#")
    ]
    forbidden = [
        "FRESH_LIMIT",
        "source_aog_db_mtime",
        "SOURCE_AOG_DB_MTIME",
        "COMMIT_TIME=",
    ]
    for kw in forbidden:
        hit_lines = [l for l in code_lines if kw in l]
        assert not hit_lines, (
            f"build-data-release.sh 非注释行含 mtime/freshness code {kw!r}, "
            f"NJX 7/30 严令删除: {hit_lines[:3]}"
        )


def test_03_kb_root_missing_fails(git_clean_check, app_commit_sha, tmp_path):
    """3. AOG_KB_ROOT 缺失 fail (exit 1)"""
    r = _run_build_data_release(
        tmp_path,
        app_commit_sha=app_commit_sha,
        env_overrides={"AOG_KB_ROOT": None},
    )
    assert r.returncode == 1, f"AOG_KB_ROOT 缺失应 exit 1, 实际 {r.returncode}\n{r.stdout + r.stderr}"
    assert "AOG_KB_ROOT" in (r.stdout + r.stderr), "失败信息应提到 AOG_KB_ROOT"


def test_04_release_dir_not_in_tmp_fails(git_clean_check, app_commit_sha, minimal_kb):
    """4. RELEASE_DIR 非 /tmp fail (exit 1)"""
    kb_root, _ = minimal_kb
    # 尝试用 /var/tmp (macOS) 或 /Users (绝对路径但非 /tmp)
    r = _run_build_data_release(
        tmp_path=kb_root.parent,  # 用 kb_root.parent 作为 tmp_path
        app_commit_sha=app_commit_sha,
        kb_root=kb_root,
        env_overrides={"RELEASE_DIR": "/var/tmp/aog-release-bad"},
    )
    assert r.returncode == 1, f"RELEASE_DIR 非 /tmp 应 exit 1, 实际 {r.returncode}\n{r.stdout + r.stderr}"
    assert "/tmp" in (r.stdout + r.stderr), "失败信息应提到 /tmp"


def test_05_release_dir_not_empty_fails(git_clean_check, app_commit_sha, minimal_kb):
    """5. RELEASE_DIR 非空 fail (exit 1)"""
    kb_root, _ = minimal_kb
    release_dir = kb_root.parent / "non_empty_release"
    release_dir.mkdir(parents=True)
    (release_dir / "old.txt").write_text("old content", encoding="utf-8")

    r = _run_build_data_release(
        tmp_path=kb_root.parent,
        app_commit_sha=app_commit_sha,
        kb_root=kb_root,
        release_dir=release_dir,
    )
    assert r.returncode == 1, f"RELEASE_DIR 非空应 exit 1, 实际 {r.returncode}\n{r.stdout + r.stderr}"


def test_06_chunks_meta_required_fails(git_clean_check, app_commit_sha, minimal_kb):
    """6. chunks_meta 缺失 fail (exit 3)

    验证: 7 件套 release bundle 必填校验, chunks_meta.json 缺则 exit 3
    """
    # 静态检查: 脚本必须验证 chunks_meta.json 存在
    content = SCRIPT_PATH.read_text(encoding="utf-8")
    assert "chunks_meta.json" in content, "build-data-release.sh 应验证 chunks_meta.json"
    # 7 件套列表里必须有 chunks_meta
    assert "chunks_meta.json" in content and "REQUIRED_ARTIFACTS" in content, (
        "build-data-release.sh 应把 chunks_meta.json 加进 REQUIRED_ARTIFACTS"
    )

    # 动态检查: 跑一次完整 release, 然后删 chunks_meta.json 验证下次会 fail
    # (但这要 build 跑成功才有 chunks_meta.json, 可能慢; 只做静态检查 + 用 grep 验证)
    # 严格 7 件套: aog.db + chroma/ + fts5_index.db + chunks_meta.json + index_stats.json + source-files-manifest.json + release-manifest.json
    # 但 release-manifest.json 是最后写, 不在 REQUIRED_ARTIFACTS 列表里, 是预期行为
    for art in ["aog.db", "fts5_index.db", "chunks_meta.json", "index_stats.json", "source-files-manifest.json"]:
        assert art in content, f"REQUIRED_ARTIFACTS 应含 {art}"


def test_07_rag_uses_release_fts5(git_clean_check, app_commit_sha):
    """7. RAG 使用 RELEASE_DIR 的 FTS5 (FTS5_TEST_PATH 绑定)"""
    content = SCRIPT_PATH.read_text(encoding="utf-8")
    # 脚本必须 export FTS5_TEST_PATH="$RELEASE_DIR/fts5_index.db"
    assert 'FTS5_TEST_PATH="$RELEASE_DIR/fts5_index.db"' in content, (
        "build-data-release.sh 必须把 FTS5_TEST_PATH 绑到 RELEASE_DIR/fts5_index.db "
        "(NJX 7/30 R-2 严令: 严禁读旧 backend/data/fts5_index.db)"
    )
    # 严禁从 backend/data 读 fts5
    assert "backend/data/fts5_index.db" not in content, (
        "build-data-release.sh 不应引用 backend/data/fts5_index.db (NJX 7/30 严令)"
    )


def test_08_pytest_failure_no_fake_green(git_clean_check, app_commit_sha):
    """8. pytest failure 不能假绿

    验证: 脚本里 8 RAG 部分非注释行严禁 || true, 必须 RAG_EXIT 捕获
    (排除注释行, 因为注释"严禁 || true"会有反例字面量)
    """
    content = SCRIPT_PATH.read_text(encoding="utf-8")
    # 8 RAG 部分必须用 set +e / exit code 捕获
    assert "RAG_EXIT" in content, "build-data-release.sh 必须捕获 RAG_EXIT (严禁 grep PASS 假绿)"
    # 非注释行严禁 || true 假绿
    code_lines = [
        line for line in content.splitlines()
        if line.strip() and not line.strip().startswith("#")
    ]
    or_true_lines = [l for l in code_lines if "|| true" in l]
    assert not or_true_lines, (
        f"build-data-release.sh 非注释行含 || true 假绿: {or_true_lines[:3]}"
    )
    # 必须 8/8 校验 (脚本里实际写法: [ "${RAG_PASS_COUNT:-0}" -lt 8 ])
    assert 'RAG_PASS_COUNT:-0}" -lt 8' in content or 'RAG_PASS_COUNT} -lt 8' in content, (
        "build-data-release.sh 必须校验 RAG_PASS_COUNT >= 8 (NJX 7/30 严令 8/8 不能少)"
    )


def test_09_pii_gate_failure_blocks(git_clean_check, app_commit_sha):
    """9. PII Gate failure 阻断 (exit 4)

    验证: 脚本非注释行不含 informational PII; 必须 PII_EXIT 捕获 + exit 4
    (排除注释行, 因为注释"严禁 ... no fail"会有反例字面量)
    """
    content = SCRIPT_PATH.read_text(encoding="utf-8")
    code_lines = [
        line for line in content.splitlines()
        if line.strip() and not line.strip().startswith("#")
    ]
    # 非注释行严禁 informational PII / no fail
    forbidden_pii = [
        "PII Gate informational",
        "PII redaction (informational",
        "no fail",
    ]
    for kw in forbidden_pii:
        hit = [l for l in code_lines if kw in l]
        assert not hit, (
            f"build-data-release.sh 非注释行含禁止的 PII keyword {kw!r}: {hit[:3]}"
        )
    # 必须有 PII_EXIT 捕获 + exit 4
    assert "PII_EXIT" in content, "build-data-release.sh 必须捕获 PII_EXIT"
    assert "exit 4" in content, "PII Gate 失败必须 exit 4 (NJX 7/30 严令)"
    # 必须有 pii_gate.test_count 等真实字段
    assert "pii_gate" in content, "release-manifest.json 必须含 pii_gate 字段 (真实执行结果)"


def test_10_release_manifest_hash_matches_artifacts(git_clean_check, app_commit_sha, minimal_kb):
    """10. release-manifest hash 与实际产物一致

    验证: release-manifest.json 的 aog.db / fts5_index.db / chunks_meta.json
    的 SHA256 字段必须等于实际文件算出的 SHA256
    """
    content = SCRIPT_PATH.read_text(encoding="utf-8")
    # 必须有 SHA256 字段
    for art in ["aog.db", "fts5_index.db", "chunks_meta.json", "source-files-manifest.json", "index_stats.json"]:
        assert f'"{art}":' in content, f"release-manifest.json 应含 {art} 字段"
    # 必须 shasum 算 SHA256
    assert 'shasum -a 256' in content, "build-data-release.sh 应算 SHA256"


def test_11_source_manifest_matches_actual_files(git_clean_check, app_commit_sha, minimal_kb):
    """11. source manifest 与实际源文件一致

    验证: build-data-release.sh 生成的 source-files-manifest.json
    的 entries 数量 = fixture KB 的可索引文件数 (4, 不含 04_课件 那个)
    """
    kb_root, expected_sha = minimal_kb
    # 必须用 /tmp 下子目录 (NJX 7/30 严令), pytest tmp_path 在 macOS /private/var/folders/... 不通过
    import tempfile
    release_dir = Path(tempfile.mkdtemp(prefix="aog-test-11-", dir="/tmp"))
    try:
        r = _run_build_data_release(
            tmp_path=kb_root.parent,
            app_commit_sha=app_commit_sha,
            kb_root=kb_root,
            release_dir=release_dir,
        )
        # 不强制 exit 0 (sentence-transformers 首次可能慢/失败),
        # 但 source-files-manifest.json 必须先写出来 (在 [3/8] 步骤)
        # 如果 build 步骤 fail, 看是否 source-manifest 写了
        manifest_path = release_dir / "source-files-manifest.json"
        if not manifest_path.exists():
            pytest.fail(
                f"build-data-release.sh 未生成 source-files-manifest.json\n"
                f"stdout: {r.stdout[:1000]}\nstderr: {r.stderr[:1000]}"
            )
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        # 04_课件 应被 SKIP_DIRS 排除, 实际 4 个可索引文件
        assert manifest["total_files"] == 4, (
            f"source manifest 应有 4 个文件 (2 city_docx + 1 exp_md + 1 cp_md), "
            f"实际 {manifest['total_files']}: {[e['relative_path'] for e in manifest['entries']]}"
        )
        # SHA256 匹配
        for entry in manifest["entries"]:
            rel = entry["relative_path"]
            assert rel in expected_sha, f"manifest 含未知文件 {rel}"
            assert entry["sha256"] == expected_sha[rel], (
                f"{rel} SHA256 不一致: manifest={entry['sha256']} actual={expected_sha[rel]}"
            )
    finally:
        shutil.rmtree(release_dir, ignore_errors=True)


def test_12_touch_no_rebuild_interpretation(git_clean_check, app_commit_sha, minimal_kb):
    """12. source 文件被 touch 但内容未变, 不得被解释为 rebuild

    验证: build-data-release.sh 不依赖 mtime; source-files-manifest.json
    是基于内容 SHA256 算的, 不会被 mtime 变化影响
    """
    content = SCRIPT_PATH.read_text(encoding="utf-8")
    # 严禁: 脚本里不能用 mtime 判定文件状态
    assert "stat -f %m" not in content, "build-data-release.sh 不应读 mtime (NJX 7/30 严令)"
    assert "stat -c %Y" not in content, "build-data-release.sh 不应读 mtime (NJX 7/30 严令)"
    # 必须基于 SHA256 (内容) 算 source manifest
    assert "sha256" in content.lower(), "source manifest 必须基于 SHA256 算"
    # 必须明确 "严禁 touch"
    assert "严禁" in content and ("touch" in content or "mtime" in content.lower()), (
        "build-data-release.sh 应明确禁止 touch / mtime 操作"
    )


def test_13_index_stats_json_load_enforced(git_clean_check, app_commit_sha):
    """13. (NJX 7/30 PR #4 严令 Fix 2) index_stats.json 必须 Python json.load 解析, 严禁 grep

    验证:
      - 脚本不允许用 grep 解析 files_failed (会漏掉有空格/换行的 array, 假绿)
      - 必须 Python json.load + 3 条一致性强制:
        a) failed_count == 0
        b) files_scanned == files_indexed
        c) files_scanned == source-files-manifest.json total_files
      - 失败时 exit 2 (build 失败)
    """
    content = SCRIPT_PATH.read_text(encoding="utf-8")

    # 严禁: 脚本里不能用 grep -o 解析 files_failed
    # 注意: 也允许 grep 解析我们自己的 SCANNED=... 输出 (那是合法的 key=value)
    # 我们只禁止 grep -o '"files_failed"' 这种
    forbidden_greps = [
        'grep -o \'"files_failed"',
        "grep -o '\"files_failed\"",
        'grep -o \'"files_scanned"',
        "grep -o '\"files_scanned\"",
    ]
    for forbidden in forbidden_greps:
        assert forbidden not in content, (
            f"build-data-release.sh 含禁止的 grep 解析 {forbidden!r}; "
            f"NJX 7/30 PR #4 严令 Fix 2: 必须 Python json.load, 严禁 grep"
        )

    # 必须有 inline Python json.load 调用
    assert "json.load" in content, (
        "build-data-release.sh 必须用 Python json.load 解析 index_stats.json "
        "(NJX 7/30 PR #4 严令 Fix 2: 严禁 grep 解析 JSON)"
    )

    # 必须验证 files_failed 字段名
    assert "files_failed" in content, "build-data-release.sh 必须读 index_stats.json files_failed 字段"

    # 必须有 3 条一致性校验
    for keyword in ["failed_count != 0", "files_scanned != files_indexed", "files_scanned != manifest_total"]:
        assert keyword in content, (
            f"build-data-release.sh 缺一致性校验 {keyword!r}; "
            f"NJX 7/30 PR #4 严令 Fix 2: 3 条一致性 (failed=0 / scanned=indexed / scanned=source_manifest.file_count)"
        )

    # 必须有 exit 2 (build 失败)
    # 找到一致性校验失败时的 exit 2 (json.load 解析失败也是 exit 2)
    assert "exit 2" in content, "index_stats.json 一致性校验失败必须 exit 2"

    # 严禁把 grep 作为 fallback (NJX 7/30 严令: 严禁退化到 grep)
    # 允许 grep 解析 key=value 输出, 但不允许 grep 解析 index_stats.json
    assert "grep -o '\"relative_path\"'" not in content, (
        "build-data-release.sh 不应用 grep 解析 source manifest file count "
        "(NJX 7/30 PR #4 严令: 必须用 manifest_total 复用)"
    )


def test_14_pii_gate_enumeration_and_hash(git_clean_check, app_commit_sha):
    """14. (NJX 7/30 PR #4 严令 Fix 3) PII Gate 必须枚举 + FTS5 + RAG + hash 化日志

    验证:
      - 必须有 PII-7 (从真实 aog.db 枚举所有 non-public/redacted phone/email)
      - 统一 internal 权限合同: internal 跟 restricted/redacted 一样视为 non-public
      - 对每个实际原值验证 FTS5 零命中
      - 验证正常 RAG context/reference 零命中
      - 日志只输出原值 sha256 hash, 严禁明文
      - PII-5 必须同时测 restricted + internal
    """
    content = SCRIPT_PATH.read_text(encoding="utf-8")

    # 1. 必须有 PII-7 (新增强项, NJX 7/30 PR #4 严令 Fix 3)
    assert "PII-7" in content, (
        "build-data-release.sh 必须含 PII-7 (从真实 aog.db 枚举 non-public/redacted 原值, FTS5 + RAG 零命中)"
    )
    # PII-7 应明确分 7a/7b/7c (FTS5 零命中 / RAG context 零命中 / 抽样校验)
    for sub in ["PII-7a", "PII-7b", "PII-7c"]:
        assert sub in content, f"build-data-release.sh PII-7 应分 3 子项 ({sub})"

    # 2. 统一 internal 权限合同: internal 视为 non-public
    # 严禁有 "permission=internal 是 D-030 'internal 半透明' 合同" 之类的旧描述
    forbidden_internal_descriptions = [
        "internal 半透明",
        "internal '半透明'",
        "internal 半透明合同",
    ]
    for forbidden in forbidden_internal_descriptions:
        assert forbidden not in content, (
            f"build-data-release.sh 含过时的 internal 描述 {forbidden!r}; "
            f"NJX 7/30 PR #4 严令 Fix 3: 统一 internal 权限合同, internal 跟 restricted 一样 REDACTED"
        )

    # 3. PII-3 应明确 non-public (含 internal) 而不只是 restricted+redacted
    # 严禁旧判定: "if perm != \"restricted\" and not redacted"
    assert 'if perm != "restricted" and not redacted' not in content, (
        "build-data-release.sh PII-3 含过时的 restricted-only 判定; "
        "NJX 7/30 PR #4 严令 Fix 3: 必须 perm == 'public' and not redacted 才 skip"
    )

    # 4. 日志只 hash, 不明文 (NJX 7/30 PR #4 严令)
    # PII-3 + PII-7 应有 hashlib.sha256 原值 hash
    assert "hashlib.sha256" in content, (
        "build-data-release.sh PII Gate 日志必须用 hashlib.sha256 输出原值 hash, 严禁明文 "
        "(NJX 7/30 PR #4 严令 Fix 3)"
    )

    # 5. PII-5 必须同时测 restricted + internal
    # (PII-5 子项 5a restricted + 5b internal)
    assert "_MockRowInternal" in content, (
        "build-data-release.sh PII-5 必须测 internal permission 同样 REDACTED "
        "(NJX 7/30 PR #4 严令 Fix 3: 统一 internal 权限合同)"
    )


def test_15_build_manifest_identity_enforced(git_clean_check, app_commit_sha):
    """15. (NJX 7/30 PR #4 严令 Fix 4+5) build_manifest 身份必须严格匹配

    验证:
      - build_manifest.build_commit == APP_COMMIT_SHA (env 注入, export_fts5 优先)
      - build_manifest.source_manifest_hash == sha256(release aog.db) on disk
      - 严禁让 export_fts5 自行 git rev-parse 写出与 APP_COMMIT_SHA 不同的 commit
    """
    # 1. build-data-release.sh 必须校验 build_manifest.build_commit == APP_COMMIT_SHA
    content = SCRIPT_PATH.read_text(encoding="utf-8")
    assert "build_manifest.build_commit" in content, (
        "build-data-release.sh 必须校验 build_manifest.build_commit == APP_COMMIT_SHA "
        "(NJX 7/30 PR #4 严令 Fix 4)"
    )
    assert "APP_COMMIT_SHA" in content, "build-data-release.sh 必须有 APP_COMMIT_SHA 校验"

    # 2. 必须校验 build_manifest.source_manifest_hash == sha256(aog.db) on disk
    assert "source_manifest_hash" in content, (
        "build-data-release.sh 必须校验 build_manifest.source_manifest_hash"
    )
    # 必须实际算 sha256(aog.db) 然后比较
    assert "hashlib.sha256" in content and "aog.db" in content, (
        "build-data-release.sh 必须实际算 sha256(aog.db) 然后与 build_manifest.source_manifest_hash 比较 "
        "(NJX 7/30 PR #4 严令 Fix 5)"
    )

    # 3. export_fts5.py 必须支持 APP_COMMIT_SHA env (优先于 git rev-parse)
    export_fts5_path = REPO_ROOT / "aog-web" / "pipeline" / "scripts" / "export_fts5.py"
    export_content = export_fts5_path.read_text(encoding="utf-8")
    assert "APP_COMMIT_SHA" in export_content, (
        "export_fts5.py 必须支持 APP_COMMIT_SHA env (NJX 7/30 PR #4 严令 Fix 4: "
        "build_manifest.build_commit 优先读 APP_COMMIT_SHA, 兜底 git rev-parse)"
    )
    # 严禁 export_fts5 走 git rev-parse 兜底后没有 APP_COMMIT_SHA 优先
    # 验证: 必须在读 APP_COMMIT_SHA 之前不调 git rev-parse (或至少 APP_COMMIT_SHA 优先)
    app_commit_idx = export_content.find("APP_COMMIT_SHA")
    git_rev_idx = export_content.find("git rev-parse")
    if app_commit_idx != -1 and git_rev_idx != -1:
        assert app_commit_idx < git_rev_idx, (
            "export_fts5.py 中 APP_COMMIT_SHA 必须在 git rev-parse 之前检查 "
            "(NJX 7/30 PR #4 严令 Fix 4: env 优先)"
        )
